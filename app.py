"""矯正診断レポート 自動生成ツール
Run: streamlit run app.py
"""

import streamlit as st
import io, os, tempfile, datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Colors ───
NAVY      = (0x0D, 0x2B, 0x4E); GOLD       = (0xB8, 0x95, 0x2A)
LIGHT_BLUE= (0xE8, 0xF1, 0xF8); WHITE      = (0xFF, 0xFF, 0xFF)
DARK      = (0x22, 0x22, 0x22); GRAY       = (0x88, 0x88, 0x88)
LGRAY_BG  = (0xF2, 0xF2, 0xF2); LGRAY      = (0xCC, 0xCC, 0xCC)
STRIPE    = (0xFA, 0xFC, 0xFF); GREEN      = (0x27, 0xAE, 0x60)
AMBER     = (0xE6, 0x7E, 0x22); AMBER_BG   = (0xFD, 0xF2, 0xE9)
GOLD_BG   = (0xFE, 0xFB, 0xEE)
GREEN     = (0x1E, 0x7E, 0x4E); GREEN_BG   = (0xEA, 0xF7, 0xEF)

FONT_NAME = 'Noto Serif JP'
LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logo_rgb.jpg')

# ─── 医院情報 ───
CLINIC_DATA = {
    'さいわいデンタルクリニック　moyuk SAPPORO院': {
        'address': '〒060-0062　北海道札幌市中央区南２条西３丁目２０番地　moyuk SAPPORO 2F',
        'tel': '050-1809-4594',
        'hours': '平日・土日　10:00〜19:00　　祝日・月曜　定休',
    },
    'さいわいデンタルクリニック': {
        'address': '〒061-1278　北海道北広島市大曲幸町４丁目４−２',
        'tel': '050-1722-4682',
        'hours': '',
    },
    'さいわいデンタルクリニック札幌大曲院': {
        'address': '〒061-1278　北海道北広島市大曲幸町６丁目１　インタービレッジ大曲',
        'tel': '011-375-7653',
        'hours': '',
    },
    'さいわいデンタルクリニック新札幌院': {
        'address': '〒004-0051　北海道札幌市厚別区厚別中央１条６丁目３−５　ラ・ジェント・ステイ新さっぽろ 2F',
        'tel': '011-802-7996',
        'hours': '',
    },
}
CLINIC_LIST = list(CLINIC_DATA.keys())

# ─── 行智会 医院情報 ───
KOCHIKAI_CLINIC_DATA = {
    'ホワイトエッセンス銀座院': {
        'address': '〒104-0061　東京都中央区銀座３丁目２−１５　ギンザ・グラッセ 5F・7F',
        'tel': '0120-884-631',
        'hours': '',
    },
    'ホワイトエッセンス新宿タカシマヤタイムズスクエア': {
        'address': '〒151-0051　東京都渋谷区千駄ケ谷５丁目２４−２　新宿タカシマヤ タイムズスクエア 12F',
        'tel': '0120-708-846',
        'hours': '',
    },
    'ホワイトエッセンス池袋': {
        'address': '〒171-0021　東京都豊島区西池袋１丁目２０−１　アイタワー 5F',
        'tel': '0120-708-846',
        'hours': '',
    },
}
KOCHIKAI_CLINIC_LIST = list(KOCHIKAI_CLINIC_DATA.keys())

def hex3(t): return f'{t[0]:02X}{t[1]:02X}{t[2]:02X}'
def rgb(t):  return RGBColor(*t)

# ─── Word helpers ───
def get_tblPr(tbl):
    p = tbl.find(qn('w:tblPr'))
    if p is None:
        p = OxmlElement('w:tblPr'); tbl.insert(0, p)
    return p

def set_cell_bg(cell, color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex3(color)); tcPr.append(shd)

def set_cell_padding(cell, top=60, bottom=60, left=140, right=140):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for s, v in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        m = OxmlElement(f'w:{s}')
        m.set(qn('w:w'), str(v)); m.set(qn('w:type'), 'dxa'); tcMar.append(m)
    tcPr.append(tcMar)

def set_cell_valign(cell, val='center'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    v = OxmlElement('w:vAlign'); v.set(qn('w:val'), val); tcPr.append(v)

def set_cell_dashed_border(cell, color=LGRAY, size=8):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{b}')
        el.set(qn('w:val'),'dashed'); el.set(qn('w:sz'),str(size))
        el.set(qn('w:space'),'0');    el.set(qn('w:color'),hex3(color))
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_no_border(table):
    tbl = table._tbl; tblPr = get_tblPr(tbl)
    tb = OxmlElement('w:tblBorders')
    for b in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{b}'); el.set(qn('w:val'),'none'); tb.append(el)
    tblPr.append(tb)

def set_border(table, color=LGRAY, size=4):
    tbl = table._tbl; tblPr = get_tblPr(tbl)
    tb = OxmlElement('w:tblBorders')
    for b in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{b}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),str(size))
        el.set(qn('w:space'),'0');    el.set(qn('w:color'),hex3(color))
        tb.append(el)
    tblPr.append(tb)

def set_tw(table):
    tbl = table._tbl; tblPr = get_tblPr(tbl)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),'5000'); tblW.set(qn('w:type'),'pct'); tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'),'fixed'); tblPr.append(tblLayout)

def set_cw(table, col, cm):
    for cell in table.columns[col].cells:
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(cm*567))); tcW.set(qn('w:type'),'dxa')
        tcPr.append(tcW)

def set_row_h(row, cm):
    tr = row._tr; trPr = tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight'); h.set(qn('w:val'), str(int(cm*567))); trPr.append(h)

def ar(para, text, bold=False, size=9, color=DARK, italic=False):
    r = para.add_run(text); r.bold=bold; r.italic=italic
    r.font.size=Pt(size); r.font.color.rgb=rgb(color)
    r.font.name = FONT_NAME
    rPr = r._r.get_or_add_rPr()
    rF = rPr.find(qn('w:rFonts'))
    if rF is None:
        rF = OxmlElement('w:rFonts'); rPr.insert(0, rF)
    rF.set(qn('w:ascii'), FONT_NAME); rF.set(qn('w:hAnsi'), FONT_NAME)
    rF.set(qn('w:eastAsia'), FONT_NAME); rF.set(qn('w:cs'), FONT_NAME)
    return r

def np(doc, text='', bold=False, size=9, color=DARK,
       align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=3):
    p = doc.add_paragraph()
    p.alignment=align; p.paragraph_format.space_before=Pt(sb)
    p.paragraph_format.space_after=Pt(sa)
    if text: ar(p, text, bold=bold, size=size, color=color)
    return p

def cp(cell, idx=0):
    p = cell.paragraphs[idx]
    p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
    return p

def cap(cell, sb=3, sa=0):
    p = cell.add_paragraph()
    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    return p

def page_header(doc, title, sub_left, sub_right):
    # ── ロゴ行 ──
    # ── タイトル帯 ──
    t = doc.add_table(rows=1, cols=1)
    set_no_border(t); set_tw(t)
    c = t.cell(0,0); set_cell_bg(c, NAVY)
    set_cell_padding(c, top=80, bottom=40, left=150, right=150)
    p = cp(c); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    ar(p, title, bold=True, size=13, color=WHITE)
    t2 = doc.add_table(rows=1, cols=2)
    set_no_border(t2); set_tw(t2); set_cw(t2,0,13); set_cw(t2,1,5)
    cl=t2.cell(0,0); cr=t2.cell(0,1)
    set_cell_bg(cl,NAVY); set_cell_bg(cr,NAVY)
    set_cell_padding(cl,top=30,bottom=50,left=150)
    set_cell_padding(cr,top=30,bottom=50,right=150)
    pl=cp(cl); ar(pl, sub_left, size=8, color=WHITE)
    pr=cp(cr); pr.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    ar(pr, sub_right, size=8, color=WHITE)
    np(doc, sa=5)

def section_label(doc, text, sb=4, sa=4):
    p = np(doc, sb=sb, sa=sa)
    r = ar(p, text, bold=True, size=11, color=NAVY)
    rPr = r._r.get_or_add_rPr()
    u = OxmlElement('w:u'); u.set(qn('w:val'),'single')
    u.set(qn('w:color'),hex3(GOLD)); rPr.append(u)

def render_complaints_box(doc, complaints):
    if not complaints:
        return
    t = doc.add_table(rows=2, cols=1)
    set_no_border(t); set_tw(t)
    # タイトル行（GOLD地）
    c0 = t.cell(0, 0)
    set_cell_bg(c0, GOLD); set_cell_padding(c0, top=50, bottom=50, left=200, right=200)
    p0 = cp(c0); ar(p0, '今回のお悩みと診断結果', bold=True, size=10, color=WHITE)
    # 本文行（クリーム地）
    c1 = t.cell(1, 0)
    set_cell_bg(c1, (255, 249, 238)); set_cell_padding(c1, top=100, bottom=100, left=220, right=220)
    first = True
    for item in complaints:
        p = cp(c1) if first else cap(c1, sb=3)
        first = False
        ar(p, f'・{item}', size=9, color=DARK)
    p_close = cap(c1, sb=8)
    ar(p_close, '今回の診断では、これらのお悩みと一致する所見が確認されました。', size=9, color=NAVY, italic=True)
    np(doc, sa=6)

def render_conclusion_box(doc, conclusion_text, rec):
    t = doc.add_table(rows=2, cols=1)
    set_no_border(t); set_tw(t)
    # タイトル行（NAVY地）
    c0 = t.cell(0, 0)
    set_cell_bg(c0, NAVY); set_cell_padding(c0, top=60, bottom=60, left=200, right=200)
    p0 = cp(c0); ar(p0, '本日の診断結果', bold=True, size=11, color=WHITE)
    # 本文行（薄青地）
    c1 = t.cell(1, 0)
    set_cell_bg(c1, LIGHT_BLUE); set_cell_padding(c1, top=120, bottom=120, left=220, right=220)
    paragraphs = [s for s in conclusion_text.split('\n') if s.strip() or True]
    first = True
    for line in paragraphs:
        p = cp(c1) if first else cap(c1, sb=2)
        first = False
        if line.startswith('推奨プラン'):
            ar(p, line, bold=True, size=11, color=GOLD)
        else:
            ar(p, line, size=9, color=DARK)
    np(doc, sa=8)

def colored_block(doc, title, body, bg=LIGHT_BLUE, tc=NAVY, bc=DARK):
    t = doc.add_table(rows=1, cols=1)
    set_no_border(t); set_tw(t); c=t.cell(0,0)
    set_cell_bg(c,bg); set_cell_padding(c,top=80,bottom=80,left=180,right=180)
    p=cp(c); ar(p, title, bold=True, size=10, color=tc)
    if body:
        p2=cap(c,sb=4); ar(p2, body, size=9, color=bc)
    np(doc, sa=4)

def embed_image_in_cell(cell, img_bytes, label, max_w=8.5, max_h=3.5):
    set_cell_valign(cell)
    if img_bytes:
        try:
            from PIL import Image as PILImage
            img = PILImage.open(io.BytesIO(img_bytes))
            iw, ih = img.size
            if iw/ih > max_w/max_h:
                dw, dh = max_w, max_w*ih/iw
            else:
                dw, dh = max_h*iw/ih, max_h
            buf = io.BytesIO(); img.save(buf, format='PNG'); buf.seek(0)
            set_cell_padding(cell, top=10, bottom=10, left=10, right=10)
            p = cp(cell); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(buf, width=Cm(dw))
            return
        except Exception:
            pass
    set_cell_bg(cell, LGRAY_BG); set_cell_dashed_border(cell)
    set_cell_padding(cell, top=60, bottom=60)
    p = cp(cell); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    ar(p, label, bold=True, size=9, color=GRAY)
    p2 = cap(cell, sb=4); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    ar(p2, '（未設定）', size=8, color=(0xAA,0xAA,0xAA), italic=True)

def set_footer_logo(doc, logo_bytes):
    """全ページのフッター中央にロゴを小さく配置する。"""
    if not logo_bytes:
        return
    try:
        from PIL import Image as PILImage
        img = PILImage.open(io.BytesIO(logo_bytes))
        iw, ih = img.size
        logo_w = 2.8          # 2.8cm幅（控えめなサイズ）
        buf = io.BytesIO(); img.save(buf, format='PNG'); buf.seek(0)
        for sec in doc.sections:
            footer = sec.footer
            footer.is_linked_to_previous = False
            # フッター段落をクリアして使う
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.clear()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fp.paragraph_format.space_before = Pt(4)
            fp.paragraph_format.space_after  = Pt(0)
            fp.add_run().add_picture(buf, width=Cm(logo_w))
            buf.seek(0)   # 次のセクション用にリセット
    except Exception:
        pass

def set_footer_text(doc, text):
    """全ページのフッター中央にテキストを配置する（行智会版用）。"""
    for sec in doc.sections:
        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(4)
        fp.paragraph_format.space_after  = Pt(0)
        r = fp.add_run(text); r.bold = True
        r.font.size = Pt(10); r.font.color.rgb = rgb(NAVY)
        r.font.name = FONT_NAME
        rPr = r._r.get_or_add_rPr()
        rF = rPr.find(qn('w:rFonts'))
        if rF is None:
            rF = OxmlElement('w:rFonts'); rPr.insert(0, rF)
        rF.set(qn('w:ascii'), FONT_NAME); rF.set(qn('w:hAnsi'), FONT_NAME)
        rF.set(qn('w:eastAsia'), FONT_NAME); rF.set(qn('w:cs'), FONT_NAME)

def make_qr(url):
    try:
        import qrcode
        qr = qrcode.QRCode(version=1, box_size=10, border=2)
        qr.add_data(url); qr.make(fit=True)
        img = qr.make_image(fill_color='black', back_color='white')
        buf = io.BytesIO(); img.save(buf, format='PNG'); buf.seek(0)
        return buf.read()
    except Exception:
        return None

# ─── Library data ───

DIAG_LIST = [
    '（選択してください）',
    '骨格性クラスII（上顎前突・出っ歯）',
    '骨格性クラスIII（下顎前突・受け口）',
    'Hyperdivergent（開咬傾向・垂直骨格）',
    'Hypodivergent（過蓋咬合傾向・水平骨格）',
    '叢生（歯のガタガタ・乱杭歯）',
    '上顎前突（歯性）',
    '下顎前突（歯性）',
    '開咬（前歯が噛まない）',
    '過蓋咬合（深い噛み合わせ）',
    '正中線のズレ',
    '空隙歯列（すきっ歯）',
    '交叉咬合（クロスバイト）',
    '埋伏歯',
    'AngleⅡ級',
    'AngleⅢ級',
]

SEV_LIST = ['軽度', '中等度', '重度']

DIAG_DESC = {
    ('骨格性クラスII（上顎前突・出っ歯）','軽度'): '上顎が下顎よりやや前方に位置しています。歯の移動を中心とした治療で改善が見込めます。',
    ('骨格性クラスII（上顎前突・出っ歯）','中等度'): '上下顎の前後的なずれが中程度認められます。マウスピース矯正で大きく改善が期待できます。',
    ('骨格性クラスII（上顎前突・出っ歯）','重度'): '骨格的なずれが顕著です。SmarteeGS（骨格矯正）を含めた治療計画が必要になる可能性があります。',
    ('骨格性クラスIII（下顎前突・受け口）','軽度'): '下顎がやや前方に突出しています。歯の傾斜を調整することで改善できます。',
    ('骨格性クラスIII（下顎前突・受け口）','中等度'): '骨格的な下顎前突が中程度あります。マウスピース矯正で歯性の改善を行います。',
    ('骨格性クラスIII（下顎前突・受け口）','重度'): '骨格的な下顎前突が著明です。SmarteeGSまたは外科的矯正が適応となる可能性があります。',
    ('Hyperdivergent（開咬傾向・垂直骨格）','軽度'): '顎の垂直方向への成長傾向が認められます。早期治療で進行を抑制できます。',
    ('Hyperdivergent（開咬傾向・垂直骨格）','中等度'): '開咬傾向の骨格パターンがあります。SmarteeGSが効果的です（北海道唯一の対応クリニック）。',
    ('Hyperdivergent（開咬傾向・垂直骨格）','重度'): '顕著な垂直骨格パターンです。SmarteeGSによる骨格矯正が必要です。一般的なマウスピース矯正では対応が難しいケースです。',
    ('Hypodivergent（過蓋咬合傾向・水平骨格）','軽度'): '顎の水平方向の成長パターンがあります。過蓋咬合の改善が必要です。',
    ('Hypodivergent（過蓋咬合傾向・水平骨格）','中等度'): '水平骨格パターンにより深い噛み合わせが生じています。治療で咬合高径を適切に改善します。',
    ('Hypodivergent（過蓋咬合傾向・水平骨格）','重度'): '著明な水平骨格パターンです。顎関節への負担が大きく、早期治療が推奨されます。',
    ('叢生（歯のガタガタ・乱杭歯）','軽度'): '軽度の歯の重なりが見られます。マウスピース矯正で6〜12ヶ月程度で改善が見込めます。',
    ('叢生（歯のガタガタ・乱杭歯）','中等度'): '中等度の叢生があります。マウスピース矯正で整列し、清掃性も大きく向上します。',
    ('叢生（歯のガタガタ・乱杭歯）','重度'): '重度の叢生があり、抜歯が必要になる場合があります。クリンチェックで最適な方針を決定します。',
    ('上顎前突（歯性）','軽度'): '上の前歯がやや前方に傾いています。マウスピース矯正で口元を整えることができます。',
    ('上顎前突（歯性）','中等度'): '上の前歯の前傾が目立ちます。口腔内の機能と審美性の改善のため、矯正治療をお勧めします。',
    ('上顎前突（歯性）','重度'): '上前歯の著明な前傾があります。唇が閉じにくく、口呼吸の原因になっている可能性があります。',
    ('開咬（前歯が噛まない）','軽度'): '前歯部に軽度の開咬が認められます。マウスピース矯正で改善可能です。',
    ('開咬（前歯が噛まない）','中等度'): '前歯が噛み合っていない状態（開咬）があります。咀嚼機能と発音に影響が出ています。',
    ('開咬（前歯が噛まない）','重度'): '重度の開咬で、骨格的な問題が主因です。SmarteeGSによる骨格矯正が必要です。',
    ('過蓋咬合（深い噛み合わせ）','軽度'): '上の前歯が下の前歯にやや深く覆いかぶさっています。矯正で改善できます。',
    ('過蓋咬合（深い噛み合わせ）','中等度'): '深い噛み合わせ（過蓋咬合）があります。顎関節や下の前歯の歯肉へのダメージが懸念されます。',
    ('過蓋咬合（深い噛み合わせ）','重度'): '重度の過蓋咬合です。下顎の動きが制限され、顎関節症のリスクが高い状態です。',
    ('正中線のズレ','軽度'): '上下の歯の中心線が約1〜2mm程度ずれています。矯正で改善可能です。',
    ('正中線のズレ','中等度'): '上下の正中線に明らかなずれがあります。顔の非対称感に影響しており、矯正治療で改善を目指します。',
    ('正中線のズレ','重度'): '著明な正中線のズレがあります。骨格的な非対称が原因の可能性があります。',
    ('空隙歯列（すきっ歯）','軽度'): '軽度のすきっ歯があります。マウスピース矯正で歯を移動し、すき間を閉鎖できます。',
    ('空隙歯列（すきっ歯）','中等度'): '複数の歯の間にすき間があります。矯正治療により審美性と機能性の両方を改善できます。',
    ('空隙歯列（すきっ歯）','重度'): '広範囲なすきっ歯があります。原因の精査（小帯・舌癖等）も含めた総合的な治療計画が必要です。',
    ('交叉咬合（クロスバイト）','軽度'): '部分的に下の歯が外側に出ている状態があります。早期治療が推奨されます。',
    ('交叉咬合（クロスバイト）','中等度'): '交叉咬合（クロスバイト）があります。顎の偏位や顔の非対称につながる可能性があります。',
    ('交叉咬合（クロスバイト）','重度'): '広範囲な交叉咬合です。顎関節への影響が大きく、早急な治療開始をお勧めします。',
    ('埋伏歯','軽度'): '歯の萌出が遅れている状態です。経過観察または矯正的牽引を行います。',
    ('埋伏歯','中等度'): '歯が骨の中に埋まっている状態（埋伏歯）があります。矯正的牽引での対応を検討します。',
    ('埋伏歯','重度'): '深部に位置する埋伏歯があります。外科的処置と矯正治療の組み合わせが必要になる場合があります。',
    ('AngleⅡ級','軽度'): '上下の第一大臼歯の咬合関係がAngle Ⅱ級を示しています。矯正治療で改善可能です。',
    ('AngleⅡ級','中等度'): 'AngleⅡ級の咬合異常があります。上顎前突を伴うことが多く、包括的な矯正治療が必要です。',
    ('AngleⅡ級','重度'): '重度のAngleⅡ級です。骨格的な問題を含む場合が多く、SmarteeGSの適応を検討します。',
    ('AngleⅢ級','軽度'): '上下の第一大臼歯の咬合関係がAngle Ⅲ級を示しています。矯正治療で改善可能です。',
    ('AngleⅢ級','中等度'): 'AngleⅢ級の咬合異常があります。下顎前突を伴うことが多く、治療計画を慎重に立案します。',
    ('AngleⅢ級','重度'): '重度のAngleⅢ級です。骨格性の要因が強く、SmarteeGSまたは外科的矯正の適応を検討します。',
    ('下顎前突（歯性）','軽度'): '下の前歯がやや前方に傾いています。矯正で改善可能です。',
    ('下顎前突（歯性）','中等度'): '下前歯の前傾が見られます。咬み合わせの改善と審美回復のため矯正治療をお勧めします。',
    ('下顎前突（歯性）','重度'): '下前歯の著明な前傾があります。咀嚼機能と審美性への影響が大きい状態です。',
}

PLAN_DATA = {
    'アライナー矯正 ミニマム':      {'price':'198,500円','period':'3〜6ヶ月',  'approach':'前歯部のみ',  'jaw':'IPRのみ',           'target':'軽度・前歯向け'},
    'アライナー矯正 ライト':        {'price':'440,000円','period':'6〜12ヶ月', 'approach':'前歯〜全歯列', 'jaw':'IPR',               'target':'軽〜中等度向け'},
    'アライナー矯正 モデレート':    {'price':'660,000円','period':'12〜18ヶ月','approach':'全歯列',      'jaw':'IPR・拡大',         'target':'中〜重度向け'},
    'アライナー矯正 コンプリヘンシブ':{'price':'990,000円','period':'18〜36ヶ月','approach':'全顎・全歯',  'jaw':'IPR・抜歯（症例による）','target':'重症例向け'},
    'SmarteeGS（骨格矯正）':        {'price':'別途見積', 'period':'要相談',   'approach':'骨格から改善', 'jaw':'骨格拡大・IPR',     'target':'骨格改善希望者'},
    'プレミアムプラン':             {'price':'899,500円〜','period':'12〜36ヶ月','approach':'矯正＋ホワイトニング','jaw':'IPR・抜歯（症例による）','target':'総合審美改善'},
    'ホワイトニング同時スタート':   {'price':'298,500円〜','period':'6〜12ヶ月','approach':'前歯＋ホワイトニング','jaw':'IPR',            'target':'審美重視の方'},
    'ワイヤー矯正':                 {'price':'800,000円〜','period':'24〜36ヶ月','approach':'全部の歯',    'jaw':'抜歯を伴う',        'target':'重症例向け'},
}

# 行智会バージョンの価格変更（モデレート・コンプリヘンシブのみ）
KOCHIKAI_PLAN_PRICES = {
    'アライナー矯正 モデレート':       '770,000円',
    'アライナー矯正 コンプリヘンシブ': '1,100,000円',
}

RISK_LIST = [
    'A: 虫歯・歯周病になりやすい\n歯並びが乱れていると歯ブラシが届きにくい部分が増え、汚れが溜まりやすくなります。その結果、虫歯や歯周病が進行しやすい状態が続きます。',
    'B: 顎（あご）が痛くなったり、口が開けにくくなる\n咬み合わせがずれていると、顎の関節に余分な負担がかかります。放置すると、食事中に顎が痛む・口を大きく開けられないなどの症状につながることがあります。',
    'C: 食べ物をうまく噛めず、胃腸に負担がかかる\n前歯や奥歯が正しく噛み合っていないと、食べ物を十分に噛み砕けません。消化不良を起こしやすくなり、体全体の健康にも影響します。',
    'D: 発音しにくい音が出てしまう\n歯の隙間や並び方によって、「さ行」「た行」などの音が発音しにくくなることがあります。仕事やコミュニケーションの場で気になる方も多いです。',
    'E: 笑顔や横顔のバランスが気になる\n歯並びは口元の印象に直結します。口を開けることへの抵抗感や、写真を撮るときに笑顔になりにくいといった悩みにつながりやすいです。',
    'F: 歯が削れたり、欠けたりしやすい\n深く噛みすぎている（過蓋咬合）や横にずれた噛み合わせ（交叉咬合）では、特定の歯に強い力が集中します。長年続くと歯が摩耗したり、割れるリスクが高まります。',
    'G: 子どもの顎や顔の成長に悪影響が出ることがある\n成長期は顎の骨が発達する大切な時期です。この時期に歯並びや噛み合わせを放置すると、顎の形や顔のバランスに影響が出ることがあります。早期対応がとくに重要です。',
]

COMMENT_PRESETS = {
    '骨格性クラスII（上顎前突・出っ歯）': [
        '今回の診断では、上下顎の前後的なずれが主な原因として認められました。骨格に由来する出っ歯は、見た目の問題にとどまらず、口が閉じにくいことによる口腔乾燥・虫歯リスクの増加、さらには顎関節への慢性的な負担にもつながります。現在の状態であれば外科手術を必要とせず、マウスピース矯正で大きく改善できる可能性があります。早く始めるほど治療期間は短くなり、ゴールも高くなります。まずはクリンチェックで3Dシミュレーションを確認し、一緒にゴールのイメージを持ちましょう。',
        '横顔のラインや口元の印象は、骨格のバランスに大きく左右されます。今回のケースは、上下顎のバランスを整えることで横顔・正面どちらも根本から改善が期待できます。放置すると骨格的なずれが徐々に進行し、顎関節症や歯の摩耗を招く可能性が高まります。矯正治療は「見た目を変える」だけでなく、「歯と顎を長く守る」ための医療です。今が最も効率よく、最小限の負担で改善できるタイミングです。ぜひ前向きにご検討ください。',
        '出っ歯は年々少しずつ悪化することが多く、放置するほど治療の選択肢が狭まります。今の段階であればマウスピース矯正のみで対応できる可能性が高く、手術なしで理想の咬み合わせと口元を手に入れられます。1000件以上の矯正経験から断言できますが、治療を始めた患者さんで「もっと早くやればよかった」とおっしゃる方が本当に多いです。この資料をご家族と共有いただき、ぜひ一緒にゴールを描いてください。',
    ],
    '骨格性クラスIII（下顎前突・受け口）': [
        '受け口（下顎前突）は骨格的な問題を含むケースが多く、放置すると年々状態が進行していきます。今回の診断では、下顎の前突が咬み合わせ・審美性・食事機能のすべてに影響していることが確認できました。当院では北海道唯一のSmarteeGSを用いることで、外科手術なしに骨格から根本改善できる可能性があります。成長が落ち着いた今のタイミングが、骨格的アプローチとして最も適した時期です。まずは無料でクリンチェックを作成し、どこまで改善できるか一緒に確認しましょう。',
        '受け口の方は前歯でうまく噛み切れないため、食事の際に奥歯だけに負担が集中しやすく、消化にも影響が出やすい状態です。また、見た目のコンプレックスから笑顔を抑えてしまっている方も多くいらっしゃいます。矯正治療によって咬み合わせを正しく整えることは、毎日の食事・会話・笑顔の質を変えることに直結します。治療を始めた多くの患者さんが「こんなに変わるとは思わなかった」とおっしゃいます。今が最善のタイミングです。',
        '下顎前突は、時間が経てば経つほど骨格の変化が固定されていきます。今の段階であれば選択できる治療法の幅が広く、最も理想に近いゴールを描けます。SmarteeGSは一般的なマウスピース矯正では対応が難しい骨格ケースにも対応できる、当院が北海道で唯一提供している治療法です。この機会に、まずはシミュレーションで未来の自分の歯並びを見てみてください。百聞は一見に如かず、必ずご納得いただけると思います。',
    ],
    'Hypodivergent（過蓋咬合傾向・水平骨格）': [
        '今回の診断では、骨格が水平方向に発達するパターン（Hypodivergent）により、深い噛み合わせ（過蓋咬合）が生じていることが確認されました。このタイプは下顎の動きが制限されやすく、顎関節に慢性的な負担がかかり続けます。放置すると歯の摩耗・顎関節症・歯肉退縮（歯茎が下がること）のリスクが高まります。今の段階であれば矯正治療で咬合高径を適切に改善できる可能性があります。早期に手を打つことが、将来の大きなトラブルを防ぐ最善策です。',
        '深い噛み合わせは「特に困っていない」と感じる方も多いのですが、実は長年にわたって歯と顎に静かなダメージを与え続けています。下の前歯が上の歯肉に当たって傷つけているケース、奥歯だけが過度に磨耗するケースも多く見られます。矯正治療で噛み合わせを整えることは、見た目の改善はもちろん、歯を長く守ることに直結します。今が動くべきタイミングです。ご家族と相談のうえ、ぜひ前向きにお考えください。',
        '骨格的な過蓋咬合傾向は、年齢を重ねるにつれてその影響が顕著になっていきます。「顎が疲れる」「肩こりや頭痛が続く」といった全身症状と咬み合わせが関係していることも少なくありません。骨格から整えることで、口腔内だけでなく全身のバランスが改善するケースもあります。1000件以上の矯正治療経験をもとに、あなたに最適なプランをご提案します。まずはシミュレーションで確認しましょう。',
    ],
    'Hyperdivergent（開咬傾向・垂直骨格）': [
        '今回の診断では、骨格が垂直方向に発達するパターン（Hyperdivergent）により、前歯が噛み合いにくい状態（開咬傾向）が確認されました。このタイプは一般的なマウスピース矯正では対応が難しいケースですが、当院では北海道唯一のSmarteeGSにより骨格から根本的な改善が可能です。今の段階であれば矯正単独での改善が見込めますが、放置すると開咬が進行し治療の難易度が上がります。今がベストなタイミングです。',
        '前歯が噛み合わない状態が続くと、奥歯だけに過度な力がかかり続けます。その結果、奥歯の摩耗・破折リスクが年々高まり、将来的には高額な補綴治療（被せもの・インプラント等）が必要になるケースも少なくありません。矯正治療はそうした将来のリスクを先手で防ぐ、最も費用対効果の高い選択肢です。SmarteeGSで骨格ごと改善し、長期的に安定した咬み合わせを手に入れましょう。',
        '開咬傾向の骨格は、発音・食事・見た目のすべてに影響を及ぼします。「食べ物をうまく噛み切れない」「さ行・た行が発音しにくい」という悩みも、咬み合わせを整えることで大きく改善できます。当院のSmarteeGSは北海道で唯一提供している治療法で、このような難症例にも対応できる専門的なアプローチです。まずはシミュレーションでゴールを確認し、一緒に最善の計画を立てましょう。',
    ],
    '叢生（歯のガタガタ・乱杭歯）': [
        '歯並びのガタガタ（叢生）は、見た目の問題だけでなく、清掃性の低下による虫歯・歯周病リスクの増加という機能的な問題も引き起こします。重なり合った歯の間には歯ブラシも歯間ブラシも届きにくく、どれだけ丁寧に磨いても限界があります。矯正治療で歯並びを整えることは、「毎日きれいに磨ける口腔環境を作ること」であり、10年後・20年後の歯の健康を根本から守ることにつながります。歯は一生使うものです。今が始め時です。',
        '叢生の程度からみると、比較的短期間で大きな変化を実感していただけるケースです。矯正治療を終えた患者さんの多くが、「笑顔に自信が持てるようになった」「写真を撮られるのが苦にならなくなった」とおっしゃいます。口元が変わると、表情・印象・自信が変わります。その変化は仕事・人間関係・日常の質にまで影響します。矯正治療は外見の投資であり、人生の投資です。ぜひ前向きにご検討ください。',
        'ガタガタした歯並びは、歯と歯が重なった部分に歯石が溜まりやすく、歯周病が静かに進行しやすい環境です。歯周病は痛みなく進行し、気づいたときには歯を支える骨が溶けていることもあります。矯正でスペースを整えることが、歯を長く守ることに最も直結します。また、マウスピース矯正であれば取り外しができるため、食事・歯磨きも普段どおりに行えます。生活への負担を最小限に、理想の歯並びを手に入れましょう。',
    ],
    '開咬（前歯が噛まない）': [
        '前歯が噛み合わない状態（開咬）は、放置すると奥歯だけに咬合力が集中し続けます。その結果、奥歯の摩耗・破折が進行し、将来的に補綴治療（被せもの・インプラント）が必要になるリスクが高まります。また、食べ物を前歯で噛み切れないため、消化への影響・発音への影響も生じています。SmarteeGSは当院が北海道で唯一提供している治療法であり、この状態の根本改善に最も適したアプローチです。今が動くべきタイミングです。',
        '開咬は「噛めない」だけでなく、発音・食事・見た目のすべてに影響を与えます。とくに「さ行」「た行」の発音に影響が出やすく、会話の中で気になっている方も多いです。しっかり前歯で噛めるようになることで、食事の楽しさ・会話の自信・笑顔の印象がまったく変わります。治療を終えた患者さんが一番おっしゃるのは「もっと早くやればよかった」という言葉です。まずはシミュレーションでゴールを確認しましょう。',
        '開咬には舌の癖（舌癖）が関係していることが多く、放置すると矯正をしても後戻りしやすい状態が続きます。当院では矯正治療と並行して舌のポジションを整えるアドバイスも行い、長期的に安定した結果を維持できるようサポートします。治療後に後戻りしにくい環境を作ることまでを含めて、真の意味での「矯正治療」と考えています。ご不安なことがあれば何でもご相談ください。',
    ],
    '過蓋咬合（深い噛み合わせ）': [
        '深い噛み合わせ（過蓋咬合）は、一見すると「よく噛めそう」に思われますが、実際には下顎の動きを制限し、顎関節に慢性的な負担をかけ続けています。放置すると顎関節症・歯肉退縮・歯の摩耗が進行し、将来的に大きな問題につながる可能性があります。矯正治療で咬合高径（噛み合わせの高さ）を適切に改善することで、顎への負担を根本から取り除けます。今が手を打つべき最善のタイミングです。',
        '過蓋咬合は年齢を重ねるほど悪化しやすく、歯や歯肉へのダメージが蓄積していきます。下の前歯が上の歯肉に当たって傷つけているケースも多く、歯茎が少しずつ下がっていく「歯肉退縮」が静かに進行していることもあります。早期に矯正治療で噛み合わせを整えることで、こうしたダメージの進行を止め、歯と歯肉を長く守ることができます。まずはシミュレーションで改善後のゴールを確認しましょう。',
        '「顎が疲れる」「肩こり・頭痛がなかなか改善しない」という方の中には、過蓋咬合による顎関節への負担が原因になっているケースがあります。咬み合わせを整えることで、口腔内だけでなく全身のバランスが改善し、こうした不調が軽減するケースも少なくありません。1000件以上の矯正治療の経験をもとに、あなたに最適なプランをご提案します。ご家族と相談のうえ、ぜひ前向きにお考えください。',
    ],
    '交叉咬合（クロスバイト）': [
        '交叉咬合（クロスバイト）は、上下の歯が本来とは逆の位置関係で噛み合っている状態です。顎に左右非対称の力がかかり続けるため、顔のゆがみ・顎関節症・特定の歯への過度な負担を引き起こします。この状態は自然に治ることはなく、放置すると年々骨格への影響が大きくなります。今のタイミングで治療を始めることが、最短・最小限の負担で改善できる道です。まずはクリンチェックでゴールを確認しましょう。',
        'クロスバイトがある場合、特定の歯に強い咬合力が集中し続けます。その結果、歯が欠けたり割れたりするリスクが高まり、将来的に補綴治療が必要になる可能性があります。矯正で正しい咬み合わせに整えることで、歯全体に力が均等に分散され、歯の寿命を大きく延ばすことができます。治療の難易度は今の段階が最も低く、選択肢も広い状態です。ぜひ前向きにご検討ください。',
        '交叉咬合は顔のゆがみとして表れることが多く、左右非対称な印象を与えることがあります。治療によって正しい咬み合わせが獲得されると、顔のバランスが整い、笑顔の印象も大きく変わります。矯正治療は機能回復と審美改善を同時に叶えられる、費用対効果の高い医療です。まずは一歩踏み出すこと。その判断が、10年後の自分を大きく変えます。',
    ],
    '正中線のずれ': [
        '上下の歯の中心線（正中線）がずれている状態が確認されました。見た目の非対称感だけでなく、顎関節への偏った負担・片側の歯への過度な咬合力という機能的な問題も生じています。矯正治療で正中を合わせることで、バランスの取れた咬み合わせと笑顔を取り戻せます。早期に対応するほど改善の幅は大きく、治療期間も短くなります。',
        '正中線のずれは、食事の際に一方の顎ばかりを使う習慣（片側咀嚼）と関係していることがあります。片側咀嚼が続くと、顎の筋肉・骨格の左右バランスが崩れ、顔のゆがみが進行します。矯正と合わせて咬む習慣を整えることで、より良い結果が長く続きます。まずはクリンチェックでゴールを確認し、改善後の自分の笑顔を見てください。',
        '正中のずれは程度によっては比較的短期間で大きく改善できるケースも多いです。口元が左右対称に整うことで、笑顔の印象・顔全体のバランスが劇的に変わります。「これほど変わるとは思わなかった」という声を多くいただく項目のひとつです。ぜひ一度、シミュレーションでゴールのイメージを確認してみてください。',
    ],
    'default': [
        '今回の診断を通じて、現在の口腔内の状態と、それが将来にどのような影響を与えるかをお伝えしました。今の状態を放置すると、歯並び・咬み合わせ・歯の健康状態は年々悪化していく可能性があります。矯正治療は「今より良くなる」だけでなく、「これ以上悪化させない」という将来への先手でもあります。1000件以上の矯正経験をもとに、あなたに最も適したプランをご提案しています。ぜひこの機会に、前向きにご検討ください。',
        '矯正治療はゴールのイメージが持てると、一気に前向きになれます。クリンチェックによる3Dシミュレーションでは、治療後の歯並びを画面上でリアルに確認できます。「どんな歯並びになるのか」「どのくらいの期間かかるのか」を見てから決めていただいて構いません。まずはシミュレーションを体験してみてください。見てから決断するのが、最善の選択です。',
        '歯並びが整うことで、清掃性・咬み合わせ機能・審美性のすべてが向上します。矯正治療のメリットは見た目の変化だけではありません。毎日の歯磨きで汚れが取りやすくなり、虫歯・歯周病リスクが下がり、食事をしっかり噛めるようになり、笑顔に自信が生まれます。これらすべての変化が、10年後・20年後の人生の質に直結します。今が決断のタイミングです。この資料をご家族と共有し、一緒に考えてみてください。',
        'マウスピース矯正は取り外し式のため、食事・歯磨きは普段どおりに行えます。見た目も目立ちにくく、仕事・プライベートで矯正装置を意識させることがほとんどありません。負担が少ない治療法だからこそ、始めやすく、続けやすいです。「やろうと思った瞬間が最速」です。今日この診断を受けたことが、最初の一歩になることを願っています。',
    ],
}

CASES_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'cases')
os.makedirs(CASES_FOLDER, exist_ok=True)

# ─── 小児矯正専用データ ───
PEDO_FINDINGS_LIST = [
    'スペース不足（永久歯が並ぶスペースが足りない）',
    '上顎前突傾向（AngleⅡ級・出っ歯傾向）',
    '下顎前突傾向（AngleⅢ級・受け口傾向）',
    '口呼吸（口で呼吸する習慣がある）',
    '口唇閉鎖力の低下（唇を閉じる力が弱い）',
    '舌突出癖（舌を前に押し出す癖がある）',
    '交叉咬合傾向（上下の歯の位置関係が一部逆になっている）',
    '過蓋咬合傾向（噛み合わせが深すぎる）',
    '開咬傾向（前歯が噛み合いにくい）',
    '叢生傾向（前歯の歯並びが乱れている）',
    '上顎歯列の狭窄（上顎の幅が狭い）',
    '下顎の成長不足傾向',
    '姿勢・口腔機能の問題（口がポカンと開く・猫背）',
]

PEDO_RISK_LIST = [
    'A: 永久歯が並ぶスペースが不足し、歯並びが乱れやすくなります\n成長が進むにつれてスペース不足が顕著になり、永久歯が重なった状態（叢生）が強くなる可能性があります。',
    'B: 出っ歯・受け口傾向が骨格ごと固定される可能性があります\n成長期に咬み合わせのズレを放置すると、顎の骨格ごと固定されていきます。成長が止まった後では骨格から改善することが困難になります。',
    'C: 将来の矯正治療で抜歯が必要になる可能性が高まります\n顎の幅が不足した状態が続くと、永久歯を並べるためのスペースを抜歯で確保する必要が出てくる可能性があります。',
    'D: 顔貌・横顔のバランス改善が難しくなります\n下顎の後退感など骨格に関わる横顔の問題は、成長期にしか改善できません。成長が終わった後では顔貌改善の選択肢が大きく制限されます。',
    'E: 口呼吸が続き、歯並びや顎の成長に悪影響を与え続けます\n口が開いた状態が続くと、舌の位置や唇の力のバランスが崩れ、歯並びや顎の成長方向に持続的な悪影響を与えます。',
    'F: 将来の治療難易度が上がり、費用・期間が増加する可能性があります\n今なら顎の成長を利用して改善できる問題も、成長後は歯を動かすだけの治療に限定されます。その結果、治療の難易度と費用・期間が増える可能性があります。',
]

PLAN_COMMENT_PRESETS = [
    {
        'label': 'A',
        'title': '軽度叢生・前歯中心パターン',
        'body': (
            '〇〇様の場合、前歯を中心に軽度の歯並びの乱れが認められます。\n\n'
            '現時点では大きな咬み合わせのズレは少なく、全体的な大きな移動までは必要ないと考えられます。そのため、できるだけ治療の負担を抑えながら、見た目と清掃性の改善を目指すプランが適していると判断しました。\n\n'
            '今回は、前歯部の歯並びを中心に整えることを目的として、「アライナー矯正 ミニマム／ライト」を推奨します。\n\n'
            '早い段階で整えておくことで、将来的な虫歯・歯周病リスクの軽減や、笑顔への自信にもつながりやすいと考えています。ご家族とも相談しながら、無理のない形で治療をご検討ください。'
        ),
    },
    {
        'label': 'B',
        'title': '中等度叢生・全体バランス改善パターン',
        'body': (
            '〇〇様の場合、前歯の歯並びだけでなく、歯列全体のバランスにも改善が必要な状態です。\n\n'
            'ライトプランでも前歯の見た目はある程度改善できる可能性がありますが、奥歯の咬み合わせや歯列全体の安定性まで考えると、治療範囲が不足する可能性があります。一方で、現時点ではコンプリヘンシブほどの大きな全顎的治療までは必要ないと判断しました。\n\n'
            'そのため、見た目・咬み合わせ・将来的な安定性のバランスを考え、「アライナー矯正 モデレート」を推奨します。\n\n'
            '矯正治療は、歯をきれいに並べるだけでなく、長く健康な状態を保つための治療でもあります。ご不安な点があれば、ご家族とも相談しながら一緒に確認していきましょう。'
        ),
    },
    {
        'label': 'C',
        'title': '重度叢生・大きなスペース不足パターン',
        'body': (
            '〇〇様の場合、歯の重なりが強く、歯を並べるためのスペースが不足している状態です。\n\n'
            '部分的な治療や短期間のプランでは、見た目の一部は改善できても、咬み合わせや仕上がりの安定性に限界が出る可能性があります。また、無理に歯を並べようとすると、歯ぐきや歯の位置に負担がかかることもあります。\n\n'
            'そのため、全体の歯並び・咬み合わせ・スペースコントロールを総合的に考え、「アライナー矯正 コンプリヘンシブ」を推奨します。\n\n'
            '治療期間はやや長くなる可能性がありますが、その分、将来的な安定性を重視した治療計画を立てることができます。ご家族とも十分に相談しながら、納得できる形で進めていきましょう。'
        ),
    },
    {
        'label': 'D',
        'title': '出っ歯・口元突出パターン',
        'body': (
            '〇〇様の場合、前歯の位置や口元の突出感が治療上のポイントになります。\n\n'
            '前歯だけをきれいに並べる治療では、口元のバランスや横顔の印象まで十分に改善できない可能性があります。そのため、歯の傾き・前後的な位置・咬み合わせを含めて、全体的に設計する必要があります。\n\n'
            '今回は、口元のバランスと咬み合わせの改善を目的として、「アライナー矯正 モデレート／コンプリヘンシブ」を推奨します。\n\n'
            '見た目の改善だけでなく、唇の閉じやすさや横顔の印象にも関わる大切な治療です。ご家族とも相談しながら、どこまでの改善を目指すかを一緒に確認していきましょう。'
        ),
    },
    {
        'label': 'E',
        'title': '受け口・反対咬合パターン',
        'body': (
            '〇〇様の場合、上下の前歯の咬み合わせに反対咬合の傾向が認められます。\n\n'
            '受け口の治療では、歯並びだけでなく、上下の顎のバランスや前後的な咬み合わせを慎重に確認する必要があります。部分的な治療では、見た目の改善に限界が出たり、咬み合わせが不安定になる可能性があります。\n\n'
            'そのため、歯列全体と骨格的なバランスを考慮し、「アライナー矯正 コンプリヘンシブ／SmarteeGS」を推奨します。\n\n'
            '必要に応じて、顎の位置や成長方向まで確認しながら治療計画を立てていきます。適応が難しい場合には無理に治療を進めず、適切な方法をご提案しますのでご安心ください。'
        ),
    },
    {
        'label': 'F',
        'title': '過蓋咬合・深い噛み合わせパターン',
        'body': (
            '〇〇様の場合、前歯の咬み込みが深く、噛み合わせの深さが治療上の重要なポイントになります。\n\n'
            '過蓋咬合は見た目だけでは気づきにくいこともありますが、前歯への負担、歯の摩耗、顎関節への負担につながることがあります。前歯だけを並べる治療では、深い咬み合わせそのものが残ってしまう可能性があります。\n\n'
            'そのため、歯並びだけでなく咬み合わせの深さまで改善する目的で、「アライナー矯正 モデレート／コンプリヘンシブ」を推奨します。\n\n'
            '将来的に歯を守るためにも、噛み合わせのバランスを整えることは大切です。ご家族とも相談しながら、見た目と機能の両面から治療をご検討ください。'
        ),
    },
    {
        'label': 'G',
        'title': '開咬・前歯が噛めていないパターン',
        'body': (
            '〇〇様の場合、前歯が十分に噛み合っておらず、開咬の傾向が認められます。\n\n'
            '開咬は、見た目だけでなく、前歯で食べ物を噛み切りにくい、発音しにくい、奥歯に負担が集中しやすいといった問題につながることがあります。前歯だけを並べる治療では、噛める状態まで改善することが難しい場合があります。\n\n'
            'そのため、歯列全体の動きと咬み合わせの安定性を考え、「アライナー矯正 コンプリヘンシブ／SmarteeGS」を推奨します。\n\n'
            '治療では、歯をきれいに並べることだけでなく、しっかり噛める状態を目指します。必要に応じて、舌の癖や口腔習癖についても確認しながら進めていきます。'
        ),
    },
    {
        'label': 'H',
        'title': 'すきっ歯・空隙歯列パターン',
        'body': (
            '〇〇様の場合、歯と歯の間に隙間があり、空隙歯列の状態が認められます。\n\n'
            '隙間を閉じるだけであれば比較的シンプルに見えることもありますが、前歯の角度や咬み合わせ、上下の歯のバランスを考えずに閉じると、仕上がりや後戻りに影響する場合があります。\n\n'
            'そのため、隙間の量と咬み合わせの状態を踏まえ、「アライナー矯正 ライト／モデレート」を推奨します。\n\n'
            '見た目の改善だけでなく、治療後の安定性まで考えて計画を立てることが大切です。ご希望の仕上がりを確認しながら、自然で美しい口元を目指していきましょう。'
        ),
    },
    {
        'label': 'I',
        'title': '後戻り・再矯正パターン',
        'body': (
            '〇〇様の場合、過去の矯正治療後に歯並びの後戻りが起きている状態です。\n\n'
            '一度矯正を経験されているため、初回治療よりも少ない範囲で改善できる可能性があります。ただし、後戻りの原因が保定だけでなく、咬み合わせや歯列全体のバランスにある場合は、部分的な治療だけでは再び後戻りする可能性があります。\n\n'
            'そのため、現在の歯並びと咬み合わせを確認したうえで、「アライナー矯正 ライト／モデレート」を推奨します。\n\n'
            '今回は、見た目の回復だけでなく、治療後に安定しやすい状態を目指すことが大切です。保定管理まで含めて、長く良い状態を維持できるようにサポートしていきます。'
        ),
    },
    {
        'label': 'J',
        'title': '顎関節・骨格改善重視パターン',
        'body': (
            '〇〇様の場合、歯並びだけでなく、顎の位置や咬み合わせ全体のバランスを考慮する必要があります。\n\n'
            '顎関節に負担がかかっている可能性があるケースでは、前歯の見た目だけを整える治療では根本的な改善につながらない場合があります。顎の位置、奥歯の咬み合わせ、上下の骨格的な関係まで含めて診断することが重要です。\n\n'
            'そのため、歯列全体と顎関節への負担を考慮し、「アライナー矯正 コンプリヘンシブ／SmarteeGS」を推奨します。\n\n'
            '治療では、歯をきれいに並べるだけでなく、顎の動きや噛み合わせの安定性も確認しながら進めていきます。必要に応じて、より精密な診断を行い、無理のない治療計画をご提案します。'
        ),
    },
]

# ─── Document generation ───

def generate_report(d):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin=Cm(1.2); sec.bottom_margin=Cm(1.2)
        sec.left_margin=Cm(1.5); sec.right_margin=Cm(1.5)

    # フォント設定（日本語含む全文字に適用）
    nml = doc.styles['Normal']
    nml.font.name = FONT_NAME
    nml.font.size = Pt(9)
    nml_rPr = nml.element.get_or_add_rPr()
    nml_rF = nml_rPr.find(qn('w:rFonts'))
    if nml_rF is None:
        nml_rF = OxmlElement('w:rFonts'); nml_rPr.insert(0, nml_rF)
    nml_rF.set(qn('w:ascii'), FONT_NAME); nml_rF.set(qn('w:hAnsi'), FONT_NAME)
    nml_rF.set(qn('w:eastAsia'), FONT_NAME); nml_rF.set(qn('w:cs'), FONT_NAME)

    # ロゴ読み込み
    logo_bytes = None
    if os.path.exists(LOGO_PATH):
        with open(LOGO_PATH, 'rb') as f:
            logo_bytes = f.read()

    pname    = d['patient_name'] or '患者名未入力'
    dstr     = d['date_str']
    diaglist = d['diagnoses']
    risks    = d['risks']
    plans    = d['plans']
    rec      = d['recommended_plan']
    pretx    = d['pretreatment'] or '特になし'
    comment  = d['dr_comment']
    conclusion = d.get('conclusion', '')
    chief_complaints = d.get('chief_complaints', [])
    expiry   = d['expiry_str']
    ctype    = d['case_type'] or '症例'
    imgs     = d['images']
    photo_comments = d.get('photo_comments', {})
    sim_url  = d['simulation_url']
    brand    = d.get('brand', 'saiwai')
    org_name = '医療法人 mirai' if brand == 'kochikai' else '医療法人 mirai　さいわいデンタルクリニック'
    plan_data_eff = {k: dict(v) for k, v in PLAN_DATA.items()}
    if brand == 'kochikai':
        for _k, _p in KOCHIKAI_PLAN_PRICES.items():
            if _k in plan_data_eff: plan_data_eff[_k]['price'] = _p

    # ── PAGE 1 ──
    page_header(doc, f'{pname}様　専用　矯正診断レポート',
                '総括責任者：理事長　谷口正昭', dstr)

    # ① 現在の口腔内の状態
    section_label(doc, '① 現在の口腔内の状態', sb=2, sa=4)
    _oral_items = [('口腔内写真','oral_photo'), ('スキャンデータ画像','scan_data')]
    t_oral = doc.add_table(rows=2, cols=2)
    set_no_border(t_oral); set_tw(t_oral)
    set_cw(t_oral,0,9); set_cw(t_oral,1,9)
    set_row_h(t_oral.rows[0], 3.2)
    for j, (lbl, key) in enumerate(_oral_items):
        embed_image_in_cell(t_oral.cell(0, j), imgs.get(key), lbl, max_w=8.5, max_h=3.2)
        # コメント行
        cc = t_oral.cell(1, j)
        comment_text = photo_comments.get(key, '')
        set_cell_bg(cc, LIGHT_BLUE)
        set_cell_padding(cc, top=60, bottom=60, left=120, right=120)
        pc = cp(cc)
        if comment_text:
            ar(pc, comment_text, size=8, color=NAVY, italic=True)
        else:
            ar(pc, ' ', size=8, color=NAVY)
    np(doc, sa=6)

    # ② 今回のお悩みと診断結果
    render_complaints_box(doc, chief_complaints)

    # ③ 本日の診断結果（結論ボックス）
    if conclusion:
        render_conclusion_box(doc, conclusion, rec)

    # ④ このまま放置すると起こり得ること
    if risks:
        t_rh=doc.add_table(rows=1,cols=1); set_no_border(t_rh); set_tw(t_rh)
        ch=t_rh.cell(0,0); set_cell_bg(ch,AMBER)
        set_cell_padding(ch,top=60,bottom=60,left=180)
        ar(cp(ch),'⚠　このまま放置すると起こり得ること',bold=True,size=10,color=WHITE)
        np(doc,sa=0)
        t_rb=doc.add_table(rows=len(risks),cols=1); set_no_border(t_rb); set_tw(t_rb)
        for i,risk in enumerate(risks):
            cell=t_rb.cell(i,0); set_cell_bg(cell,AMBER_BG)
            set_cell_padding(cell,top=50,bottom=50,left=200)
            ar(cp(cell),f'▶　{risk}',size=9,color=AMBER)
        np(doc,sa=4)

    # ⑤ 治療によって期待できる変化
    benefits = d.get('benefits', [])
    if benefits:
        t_bh=doc.add_table(rows=1,cols=1); set_no_border(t_bh); set_tw(t_bh)
        bh=t_bh.cell(0,0); set_cell_bg(bh,GREEN)
        set_cell_padding(bh,top=60,bottom=60,left=180)
        ar(cp(bh),'✓　治療によって期待できる変化',bold=True,size=10,color=WHITE)
        np(doc,sa=0)
        t_bb=doc.add_table(rows=len(benefits),cols=1); set_no_border(t_bb); set_tw(t_bb)
        for i,benefit in enumerate(benefits):
            cell=t_bb.cell(i,0); set_cell_bg(cell,GREEN_BG)
            set_cell_padding(cell,top=50,bottom=50,left=200)
            ar(cp(cell),f'◎　{benefit}',size=9,color=GREEN)
    np(doc,sa=6); doc.add_page_break()

    # ── PAGE 1.5（診断詳細）── ページ2冒頭に移動
    page_header(doc, f'{pname}様　専用　矯正診断レポート',
                '総括責任者：理事長　谷口正昭', dstr)
    section_label(doc, '今日の検査でわかったこと')
    nums = '①②③④'
    active = [g for g in diaglist if g.get('item') and g['item'] != '（選択してください）']
    for i, diag in enumerate(active[:4]):
        t = doc.add_table(rows=1, cols=1)
        set_no_border(t); set_tw(t); c=t.cell(0,0)
        set_cell_bg(c, LIGHT_BLUE)
        set_cell_padding(c, top=70, bottom=70, left=180, right=180)
        p=cp(c)
        ar(p, f'{nums[i]}  ', bold=True, size=11, color=NAVY)
        ar(p, diag['item'], bold=True, size=11, color=NAVY)
        ar(p, '　　重症度：', size=9, color=GRAY)
        ar(p, diag['severity'], bold=True, size=10, color=GOLD)
        p2=cap(c, sb=4); ar(p2, diag.get('description',''), size=9, color=DARK)
        np(doc, sa=3)
    np(doc, sa=6); doc.add_page_break()

    # ── PAGE 3（治療プラン） ──
    page_header(doc,'あなたへの治療プラン','総括責任者：理事長　谷口正昭',dstr)
    section_label(doc,'提案プラン比較')

    pcols=[3.5,4.5,4.5,5.5]
    pt=doc.add_table(rows=6,cols=4); set_border(pt,color=LGRAY,size=4); set_tw(pt)
    for j,w in enumerate(pcols): set_cw(pt,j,w)

    for j in range(4):
        c=pt.cell(0,j)
        if j==0:
            bg=LGRAY_BG; fc=DARK; txt=''
        else:
            pname_j = plans[j-1]['name'] if j-1 < len(plans) else '—'
            bg = GOLD if pname_j == rec else NAVY
            fc = WHITE; txt = pname_j
        set_cell_bg(c,bg); set_cell_padding(c,top=70,bottom=70)
        p=cp(c); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        ar(p, txt, bold=True, size=9, color=fc)

    row_labels=['治療期間','費用目安','対象範囲','スペース確保の方法','こんな方に']
    row_keys  =['period',  'price',   'approach','jaw',               'target']
    for i,(lbl,key) in enumerate(zip(row_labels,row_keys)):
        row=pt.rows[i+1]
        c0=row.cells[0]; set_cell_bg(c0,LGRAY_BG)
        set_cell_padding(c0,top=55,bottom=55,left=120)
        ar(cp(c0),lbl,bold=True,size=9,color=NAVY)
        for j in range(1,4):
            c=row.cells[j]
            pn = plans[j-1]['name'] if j-1 < len(plans) else None
            txt = plan_data_eff[pn].get(key,'—') if pn and pn in plan_data_eff else '—'
            if i%2==1: set_cell_bg(c,LIGHT_BLUE)
            set_cell_padding(c,top=55,bottom=55)
            p=cp(c); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            ar(p,txt,size=9,color=DARK)

    np(doc,sa=4)
    p_rec=np(doc,sb=2,sa=6)
    ar(p_rec,'★ Dr.推奨プラン：',bold=True,size=10,color=GOLD)
    ar(p_rec, rec, bold=True,size=10,color=NAVY)

    colored_block(doc,'矯正前に必要な処置',pretx,bg=LIGHT_BLUE)

    # QR
    qr_bytes = make_qr(sim_url) if sim_url else None
    t_qr=doc.add_table(rows=1,cols=2); set_no_border(t_qr); set_tw(t_qr)
    set_cw(t_qr,0,13); set_cw(t_qr,1,5); set_row_h(t_qr.rows[0],3.0)
    cl_qr=t_qr.cell(0,0)
    set_cell_padding(cl_qr,top=60,bottom=60,left=150); set_cell_valign(cl_qr)
    ar(cp(cl_qr),'推奨プランのシミュレーションはこちら　→',bold=True,size=9.5,color=NAVY)
    embed_image_in_cell(t_qr.cell(0,1), qr_bytes,'QRコード',max_w=4.0,max_h=3.0)
    np(doc,sa=5)

    # Dr. comment
    dc_t=doc.add_table(rows=1,cols=2); set_no_border(dc_t); set_tw(dc_t)
    set_cw(dc_t,0,0.3); set_cw(dc_t,1,17.7)
    set_cell_bg(dc_t.cell(0,0),NAVY)
    cr_b=dc_t.cell(0,1); set_cell_bg(cr_b,(0xF8,0xF9,0xFA))
    set_cell_padding(cr_b,top=90,bottom=90,left=220,right=180)
    ar(cp(cr_b),'理事長　谷口正昭より',bold=True,size=10,color=GOLD)
    p_dc=cap(cr_b,sb=6); p_dc.paragraph_format.space_after=Pt(10)
    ar(p_dc,f'「{comment}」',italic=True,size=9,color=DARK)
    p_ds=cap(cr_b); p_ds.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    ar(p_ds,f'{org_name}　理事長　谷口 正昭',bold=True,size=9,color=GOLD)
    np(doc,sa=4)

    # Payment
    exp_lbl = expiry if expiry else '——'
    pay_rows=[
        ('◎ デンタルクレジット','12回払い　金利0円'),
        ('◎ 医療費控除の対象','確定申告で一部還付可能（年間10万円超）'),
        ('→ 進める場合',f'クリンチェック申込：33,000円→0円（期間限定）　{exp_lbl}まで有効'),
    ]
    pt2=doc.add_table(rows=3,cols=2); set_border(pt2,color=LGRAY,size=4); set_tw(pt2)
    set_cw(pt2,0,5.5); set_cw(pt2,1,12.5)
    for i,(lbl,cont) in enumerate(pay_rows):
        row=pt2.rows[i]; cl2=row.cells[0]; cr2=row.cells[1]
        if i%2==0: set_cell_bg(cl2,LIGHT_BLUE); set_cell_bg(cr2,STRIPE)
        set_cell_padding(cl2,top=55,bottom=55,left=120)
        set_cell_padding(cr2,top=55,bottom=55,left=120)
        ar(cp(cl2),lbl,bold=True,size=9,color=NAVY)
        if i==2:
            base=cont.split(exp_lbl)[0]
            ar(cp(cr2),base,size=9,color=DARK)
            ar(cp(cr2) if False else cr2.paragraphs[0],exp_lbl,bold=True,size=9,color=GOLD)
            cr2.paragraphs[0].add_run('まで有効').font.size=Pt(9)
        else:
            ar(cp(cr2),cont,size=9,color=DARK)
    doc.add_page_break()

    # ── PAGE 3 ──
    _p3_title = '矯正方法の選び方・医療法人 mirai を選ぶ理由' if brand == 'kochikai' else '矯正方法の選び方・医療法人 mirai さいわいを選ぶ理由'
    page_header(doc,_p3_title,'総括責任者：理事長　谷口正昭',dstr)
    section_label(doc,f'{ctype}の治療例')

    case_t=doc.add_table(rows=2,cols=3); set_border(case_t,color=LGRAY,size=4); set_tw(case_t)
    set_cw(case_t,0,8); set_cw(case_t,1,1.5); set_cw(case_t,2,8.5)
    for i in range(2):
        set_row_h(case_t.rows[i],3.2)
        embed_image_in_cell(case_t.cell(i,0), imgs.get(f'before{i+1}'),'Before',max_w=7.5,max_h=3.2)
        c_ar=case_t.cell(i,1); set_cell_bg(c_ar,WHITE)
        set_cell_padding(c_ar,top=60,bottom=60); set_cell_valign(c_ar)
        p_ar=cp(c_ar); p_ar.alignment=WD_ALIGN_PARAGRAPH.CENTER
        ar(p_ar,'→',size=14,color=(0xBB,0xBB,0xBB),bold=True)
        embed_image_in_cell(case_t.cell(i,2), imgs.get(f'after{i+1}'),'After',max_w=8.0,max_h=3.2)
    np(doc,sa=5)

    # Comparison 1
    section_label(doc,'ワイヤー矯正　vs　マウスピース矯正')
    c1r=[('見た目','装置が目立つ','ほぼ透明・気づかれにくい ✓'),
         ('取り外し','できない','食事・歯磨き時に外せる ✓'),
         ('歯磨き','磨きにくい','通常通り可能 ✓'),
         ('通院頻度','月1〜2回','2〜3ヶ月に1回 ✓'),
         ('痛み','強め','比較的少ない ✓'),
         ('費用目安','80〜120万円','66〜150万円')]
    ct1=doc.add_table(rows=len(c1r)+1,cols=3); set_border(ct1,color=LGRAY,size=4); set_tw(ct1)
    for j,w in enumerate([3.0,6.0,9.0]): set_cw(ct1,j,w)
    for j,(txt,bg) in enumerate([('',LGRAY_BG),('ワイヤー矯正',LGRAY),('マウスピース矯正',LIGHT_BLUE)]):
        c=ct1.cell(0,j); set_cell_bg(c,bg); set_cell_padding(c,top=60,bottom=60)
        p=cp(c); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; ar(p,txt,bold=True,size=9,color=NAVY)
    for i,(lbl,wire,mp) in enumerate(c1r):
        row=ct1.rows[i+1]
        for j,(txt,align) in enumerate([(lbl,WD_ALIGN_PARAGRAPH.LEFT),(wire,WD_ALIGN_PARAGRAPH.CENTER),(mp,WD_ALIGN_PARAGRAPH.LEFT)]):
            c=row.cells[j]
            if j==0: set_cell_bg(c,LGRAY_BG)
            elif i%2==1: set_cell_bg(c,STRIPE)
            set_cell_padding(c,top=45,bottom=45,left=120)
            p=cp(c); p.alignment=align
            ar(p,txt,bold=(j==0),size=8.5,color=NAVY if j==0 else (GREEN if '✓' in txt else DARK))
    np(doc,sa=5)

    # Comparison 2
    section_label(doc,f'他院のマウスピース矯正　vs　{org_name}')
    c2r=[('診断精度','X線のみが多い','CBCT＋セファロ＋3Dスキャン ★'),
         ('骨格矯正','非対応が多い','SmarteeGS対応（北海道唯一）★'),
         ('診断レポート','有料3〜5万円','完全無料 ★'),
         ('医師の専門性','様々','世界TOP1% ダイアモンドドクター ★'),
         ('対応症例','軽〜中等度のみ多い','重症例・骨格レベルの治療まで一貫対応 ★'),
         ('適応外の場合','そのまま受入も','正直にお伝えし他院紹介 ★')]
    ct2=doc.add_table(rows=len(c2r)+1,cols=3); set_border(ct2,color=LGRAY,size=4); set_tw(ct2)
    for j,w in enumerate([3.5,6.0,8.5]): set_cw(ct2,j,w)
    for j,(txt,bg,fc) in enumerate([('',LGRAY_BG,NAVY),('一般的なクリニック',LGRAY,DARK),(org_name,NAVY,WHITE)]):
        c=ct2.cell(0,j); set_cell_bg(c,bg); set_cell_padding(c,top=60,bottom=60)
        p=cp(c); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; ar(p,txt,bold=True,size=9,color=fc)
    for i,(lbl,other,saiwai) in enumerate(c2r):
        row=ct2.rows[i+1]
        for j,(txt,align) in enumerate([(lbl,WD_ALIGN_PARAGRAPH.LEFT),(other,WD_ALIGN_PARAGRAPH.CENTER),(saiwai,WD_ALIGN_PARAGRAPH.LEFT)]):
            c=row.cells[j]
            if j==0: set_cell_bg(c,LGRAY_BG)
            elif j==2 and i%2==0: set_cell_bg(c,LIGHT_BLUE)
            elif j==1 and i%2==1: set_cell_bg(c,STRIPE)
            set_cell_padding(c,top=45,bottom=45,left=120)
            p=cp(c); p.alignment=align
            ar(p,txt,bold=(j==0),size=8.5,color=NAVY if j==0 else (GREEN if '★' in txt else DARK))
    np(doc,sa=5)

    # Dr. photo + credentials
    t_cred=doc.add_table(rows=1,cols=1); set_no_border(t_cred); set_tw(t_cred)
    cc=t_cred.cell(0,0); set_cell_bg(cc,LIGHT_BLUE)
    set_cell_padding(cc,top=80,bottom=80,left=180)
    ar(cp(cc),f'{org_name}　理事長　谷口 正昭',bold=True,size=10,color=GOLD)
    np(doc,sa=3)
    t_photo=doc.add_table(rows=1,cols=2); set_no_border(t_photo); set_tw(t_photo)
    set_cw(t_photo,0,4.5); set_cw(t_photo,1,13.5); set_row_h(t_photo.rows[0],5.0)
    embed_image_in_cell(t_photo.cell(0,0), imgs.get('dr_photo'),'理事長の写真',max_w=4.0,max_h=5.0)
    c_cr=t_photo.cell(0,1); set_cell_padding(c_cr,top=60,bottom=60,left=180); set_cell_valign(c_cr)
    first=True
    for cred in ['★　インビザライン 世界TOP1%　ダイアモンドドクター\n　　通算症例数 1,000件以上',
                 '★　SmarteeGS認定取得医\n　　日本全国でわずか数名のみ・日本初 SmarteeGS 実施 Dr',
                 '★　CBCT・セファロ・3D口腔内スキャン\n　　骨格レベルから治療を設計']:
        p_ci = cp(c_cr) if first else cap(c_cr,sb=6)
        first=False; ar(p_ci,cred,size=9,color=NAVY)
    np(doc,sa=5); doc.add_page_break()

    # ── PAGE 4 ──
    page_header(doc,'ご家族の方へ・次のステップ','総括責任者：理事長　谷口正昭',dstr)
    t_fam=doc.add_table(rows=1,cols=1); set_no_border(t_fam); set_tw(t_fam)
    cf=t_fam.cell(0,0); set_cell_bg(cf,LIGHT_BLUE)
    set_cell_padding(cf,top=100,bottom=100,left=200,right=200)
    ar(cp(cf),'このレポートをご覧のご家族の方へ',bold=True,size=11,color=NAVY)
    fam_paras = [
        '今回の矯正治療は、見た目だけを目的としたものではありません。\n歯並びや咬み合わせの乱れは、清掃性・虫歯リスク・歯周病リスク・歯への負担・顎関節への負担に関係することがあります。',
        'もちろん、矯正治療は決して安い治療ではありません。\nそのため当院では、必要性・治療方法・費用・期間を明確にしたうえで、ご家族で納得して判断していただくことを大切にしています。',
        '今回おすすめしているプランは、見た目の改善だけでなく、長期的な噛み合わせの安定性まで考えたご提案です。',
        f'{pname}さまの将来の歯を守るために、今どの選択が一番良いかを一緒に考えていただければと思います。',
    ]
    for i, para in enumerate(fam_paras):
        pf = cap(cf, sb=(8 if i == 0 else 6))
        if i < len(fam_paras) - 1:
            pf.paragraph_format.space_after = Pt(6)
        ar(pf, para, size=9, color=DARK)
    pf3=cap(cf,sb=8); pf3.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    ar(pf3,f'{org_name}　理事長　谷口 正昭',bold=True,size=8.5,color=GOLD)
    np(doc,sa=5)

    # FAQ
    section_label(doc,'よくあるご質問')

    def faq_item(q, a_lines):
        fq=doc.add_table(rows=2,cols=1); set_no_border(fq); set_tw(fq)
        cq=fq.cell(0,0); set_cell_bg(cq,(0xF0,0xF4,0xF8))
        set_cell_padding(cq,top=45,bottom=45,left=160)
        ar(cp(cq),'Q.　'+q,bold=True,size=9,color=NAVY)
        ca=fq.cell(1,0)
        tc2=ca._tc; tcPr2=tc2.get_or_add_tcPr()
        tcB=OxmlElement('w:tcBorders'); lb=OxmlElement('w:left')
        lb.set(qn('w:val'),'single'); lb.set(qn('w:sz'),'10')
        lb.set(qn('w:space'),'0'); lb.set(qn('w:color'),hex3(NAVY))
        tcB.append(lb); tcPr2.append(tcB)
        set_cell_padding(ca,top=45,bottom=55,left=180)
        for i,(line,bf,cf2) in enumerate(a_lines):
            p_a=cp(ca) if i==0 else cap(ca,sb=3)
            ar(p_a,('A.　' if i==0 else '')+line,size=9,color=cf2,bold=bf)
        np(doc,sa=3)

    # Cost sub-header
    ts1=doc.add_table(rows=1,cols=1); set_no_border(ts1); set_tw(ts1)
    cs1=ts1.cell(0,0); set_cell_bg(cs1,GOLD_BG)
    set_cell_padding(cs1,top=30,bottom=30,left=160)
    ar(cp(cs1),'費用・お支払いについて',bold=True,size=9,color=GOLD)
    np(doc,sa=2)
    faq_item('月々いくらから始められますか？',[
        ('デンタルクレジットで12回払い・金利0円でお支払いいただけます。',False,DARK),
        ('例）アライナー矯正 ミニマム（330,000円）→ 月々 27,500円〜',False,NAVY),
        ('　　アライナー矯正 ライト（440,000円）→ 月々 36,667円〜',False,NAVY),
        ('　　アライナー矯正 モデレート（660,000円）→ 月々 55,000円〜',False,NAVY),
        (f'▶  クリンチェック申込 33,000円 → 0円（期間限定・{exp_lbl}まで）',True,GOLD),
    ])
    faq_item('医療費控除は使えますか？いくら戻りますか？',[
        ('矯正治療は医療費控除の対象です（年間医療費10万円超の場合）。',False,DARK),
        ('例）年収200万円・治療費40万円 → 約30,000円還付',False,NAVY),
        ('　　年収300万円・治療費40万円 → 約60,000円還付',False,NAVY),
        ('　　年収700万円・治療費40万円 → 約90,100円還付',False,NAVY),
        ('確定申告で申請できます。領収書は必ず保管してください。',False,DARK),
    ])
    np(doc,sa=2)
    ts2=doc.add_table(rows=1,cols=1); set_no_border(ts2); set_tw(ts2)
    cs2=ts2.cell(0,0); set_cell_bg(cs2,LIGHT_BLUE)
    set_cell_padding(cs2,top=30,bottom=30,left=160)
    ar(cp(cs2),'治療・生活について',bold=True,size=9,color=NAVY)
    np(doc,sa=2)
    faq_item('治療中は痛いですか？',[
        ('マウスピース矯正はワイヤーに比べ痛みが少ないです。',False,DARK),
        ('装置交換後1〜3日は軽い圧迫感を感じる方が多いですが、慣れると気になりません。',False,DARK),
    ])
    faq_item('仕事や学校に通いながらできますか？',[
        ('はい。通院は2〜3ヶ月に1回程度です。',False,DARK),
        ('装置は透明なため外見を気にせず生活でき、接客・営業・プレゼンも問題ありません。',False,DARK),
    ])
    faq_item('どれくらいで効果が出ますか？',[
        ('症状によりますが、多くの方が6〜12ヶ月で変化を実感されています。',False,DARK),
        ('歯並びの改善は通常1〜3年程度です。骨格改善が必要な場合は別途ご相談します。',False,DARK),
    ])
    np(doc,sa=5)

    # Next steps
    t_ns=doc.add_table(rows=1,cols=1); set_no_border(t_ns); set_tw(t_ns)
    cn=t_ns.cell(0,0); set_cell_bg(cn,NAVY)
    set_cell_padding(cn,top=100,bottom=100,left=200,right=200)
    ar(cp(cn),'次のステップ',bold=True,size=11,color=WHITE)
    _steps = [
        ('①まずはご家族とこのレポートをご確認ください',
         '特に「推奨プラン」と「費用・期間」をご確認ください。'),
        ('②不安な点をLINEで送ってください',
         '費用、痛み、期間、支払い方法、治療開始時期など、どんな内容でも大丈夫です。'),
        ('③必要であれば15分の無料再相談が可能です',
         '来院またはオンラインで、ご家族同席でもご相談いただけます。'),
        ('④治療を進める場合は手続きに進みます',
         '本日ご確認いただいたシミュレーションと治療方針をもとに、治療開始に必要なお手続き・お支払い方法・開始時期をご案内いたします。'),
    ]
    for title, detail in _steps:
        pt=cap(cn,sb=8); ar(pt, title, bold=True, size=9.5, color=GOLD)
        pd=cap(cn,sb=2); ar(pd, detail, size=8.5, color=WHITE)
    cap(cn,sb=10)
    _closing = [
        '矯正治療は、すぐに決める必要はありません。',
        'ただし、歯並びや咬み合わせの問題は、自然に改善することはほとんどありません。',
        '',
        '大切なのは、今すぐ始めるかどうかではなく、',
        'ご自身の状態を正しく理解し、将来後悔しない選択をすることです。',
        '',
        '不安な点があれば、遠慮なくご相談ください。',
        '無理に治療をすすめるのではなく、納得して判断していただくことを大切にしています。',
    ]
    for line in _closing:
        pc=cap(cn,sb=2); ar(pc, line, size=8.5, color=WHITE, italic=(line==''))
    np(doc,sa=5)

    # Clinic info
    t_cl=doc.add_table(rows=1,cols=1); set_tw(t_cl)
    tblPr_el=get_tblPr(t_cl._tbl); borders_el=OxmlElement('w:tblBorders')
    for b in ['top','left','bottom','right']:
        el=OxmlElement(f'w:{b}'); el.set(qn('w:val'),'single')
        el.set(qn('w:sz'),'8'); el.set(qn('w:space'),'0')
        el.set(qn('w:color'),hex3(GOLD)); borders_el.append(el)
    for b in ['insideH','insideV']:
        el=OxmlElement(f'w:{b}'); el.set(qn('w:val'),'none'); borders_el.append(el)
    tblPr_el.append(borders_el)
    ccl=t_cl.cell(0,0); set_cell_padding(ccl,top=80,bottom=80,left=150,right=150)
    pc1=cp(ccl); pc1.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _cn = d.get('clinic_name', CLINIC_LIST[0])
    _ci = d.get('clinic_info', CLINIC_DATA[CLINIC_LIST[0]])
    ar(pc1, _cn, bold=True, size=10, color=NAVY)
    _cl = [_ci['address'], f'TEL : {_ci["tel"]}']
    if _ci.get('hours'): _cl.append(_ci['hours'])
    for line in _cl:
        pc2=cap(ccl,sb=3); pc2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        ar(pc2,line,size=8.5,color=DARK)

    if brand == 'kochikai':
        set_footer_text(doc, 'mirai')
    else:
        set_footer_logo(doc, logo_bytes)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


def generate_pedo_report(d):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin=Cm(1.2); sec.bottom_margin=Cm(1.2)
        sec.left_margin=Cm(1.5); sec.right_margin=Cm(1.5)
    nml = doc.styles['Normal']
    nml.font.name = FONT_NAME; nml.font.size = Pt(9)
    nml_rPr = nml.element.get_or_add_rPr()
    nml_rF = nml_rPr.find(qn('w:rFonts'))
    if nml_rF is None:
        nml_rF = OxmlElement('w:rFonts'); nml_rPr.insert(0, nml_rF)
    nml_rF.set(qn('w:ascii'), FONT_NAME); nml_rF.set(qn('w:hAnsi'), FONT_NAME)
    nml_rF.set(qn('w:eastAsia'), FONT_NAME); nml_rF.set(qn('w:cs'), FONT_NAME)

    logo_bytes = None
    if os.path.exists(LOGO_PATH):
        with open(LOGO_PATH, 'rb') as f:
            logo_bytes = f.read()

    child_name       = d.get('child_name', '')
    guardian_name    = d.get('guardian_name', '')
    child_age        = d.get('child_age', '')
    dstr             = d.get('date_str', '')
    findings         = d.get('findings', [])
    space_shortage   = d.get('space_shortage', '')
    rec_device       = d.get('recommended_device', 'ブランベラボキッズプログラム')
    period_1         = d.get('treatment_period_1', '約1〜2年')
    cost_1           = d.get('treatment_cost_1', '440,000円（税込）')
    monthly_fee      = d.get('monthly_fee', '5,500円/月（税込）')
    phase2_note      = d.get('phase2_note', '')
    risks            = d.get('risks_pedo', [])
    comment          = d.get('dr_comment', '')
    expiry           = d.get('expiry_str', '')
    imgs             = d.get('images', {})

    # ── PAGE 1：表紙・診断結論 ──
    page_header(doc, f'{child_name}様　小児矯正診断レポート',
                '総括責任者：理事長　谷口正昭', dstr)

    t_guard = doc.add_table(rows=1, cols=1)
    set_no_border(t_guard); set_tw(t_guard)
    cg = t_guard.cell(0, 0); set_cell_bg(cg, GOLD_BG)
    set_cell_padding(cg, top=60, bottom=60, left=180, right=180)
    ar(cp(cg), f'保護者 {guardian_name}様へ', bold=True, size=9, color=GOLD)
    pg2 = cap(cg, sb=4)
    ar(pg2, f'このレポートは、本日のカウンセリングにご参加いただけなかった保護者の方に、{child_name}さんの診断内容をわかりやすくお伝えするために作成しました。ご家族でお読みいただき、ご不明な点はLINEまたはお電話でお気軽にご相談ください。', size=8.5, color=DARK)
    np(doc, sa=5)

    t_conc = doc.add_table(rows=2, cols=1)
    set_no_border(t_conc); set_tw(t_conc)
    c_ch = t_conc.cell(0, 0); set_cell_bg(c_ch, NAVY)
    set_cell_padding(c_ch, top=60, bottom=60, left=200, right=200)
    ar(cp(c_ch), f'本日の診断結果　／　{child_name}さん（{child_age}）', bold=True, size=11, color=WHITE)
    c_cb = t_conc.cell(1, 0); set_cell_bg(c_cb, LIGHT_BLUE)
    set_cell_padding(c_cb, top=100, bottom=100, left=220, right=220)
    if findings:
        ar(cp(c_cb), '確認された所見：', bold=True, size=9, color=NAVY)
        for fi in findings:
            p_fi = cap(c_cb, sb=3); ar(p_fi, f'・{fi}', size=9, color=DARK)
    p_rec = cap(c_cb, sb=10)
    ar(p_rec, '推奨プラン：', bold=True, size=11, color=GOLD)
    ar(p_rec, rec_device, bold=True, size=11, color=NAVY)
    if phase2_note:
        p_ph = cap(c_cb, sb=6)
        ar(p_ph, f'※ {phase2_note}', size=8.5, color=GRAY, italic=True)
    np(doc, sa=5)

    t_ph = doc.add_table(rows=1, cols=4)
    set_border(t_ph, color=LGRAY, size=4); set_tw(t_ph)
    set_row_h(t_ph.rows[0], 3.5)
    for j, (lbl, key) in enumerate([('顔貌（正面）','face_front'),('顔貌（横顔）','face_side'),
                                      ('口腔内写真','oral'),('スキャンデータ','scan')]):
        embed_image_in_cell(t_ph.cell(0, j), imgs.get(key), lbl, max_w=4.0, max_h=3.5)
    np(doc, sa=3); doc.add_page_break()

    # ── PAGE 2：今日の検査でわかったこと ──
    page_header(doc, f'{child_name}様　小児矯正診断レポート',
                '総括責任者：理事長　谷口正昭', dstr)
    section_label(doc, '今日の検査でわかったこと')
    nums = '①②③④⑤⑥⑦⑧'
    for i, fi in enumerate(findings[:8]):
        t_fi = doc.add_table(rows=1, cols=1)
        set_no_border(t_fi); set_tw(t_fi)
        c_fi = t_fi.cell(0, 0); set_cell_bg(c_fi, LIGHT_BLUE)
        set_cell_padding(c_fi, top=60, bottom=60, left=180, right=180)
        p_fi = cp(c_fi)
        ar(p_fi, f'{nums[i]}  ', bold=True, size=11, color=NAVY)
        ar(p_fi, fi, bold=True, size=10, color=NAVY)
        if space_shortage and i == 0 and 'スペース不足' in fi:
            p_fi2 = cap(c_fi, sb=3)
            ar(p_fi2, f'→ 不足量の目安：{space_shortage}', size=9, color=DARK)
        np(doc, sa=3)
    np(doc, sa=3)
    colored_block(doc,
        '小児矯正では「歯並び」だけでなく、顎の成長・口呼吸・舌や唇の使い方も改善します',
        '歯並びの問題は見た目だけでなく、呼吸・噛む力・話し方・顔の成長すべてに関係しています。成長期の今だからこそ、幅広い問題にアプローチすることができます。',
        bg=GOLD_BG, tc=GOLD, bc=DARK)
    np(doc, sa=4); doc.add_page_break()

    # ── PAGE 3：放置リスク ──
    page_header(doc, f'{child_name}様　小児矯正診断レポート',
                '総括責任者：理事長　谷口正昭', dstr)
    section_label(doc, 'このまま成長を見守るだけだと起こり得ること')
    if risks:
        t_rh = doc.add_table(rows=1, cols=1); set_no_border(t_rh); set_tw(t_rh)
        rh = t_rh.cell(0, 0); set_cell_bg(rh, AMBER)
        set_cell_padding(rh, top=60, bottom=60, left=180)
        ar(cp(rh), '▶　放置した場合に起こり得ること', bold=True, size=10, color=WHITE)
        np(doc, sa=0)
        t_rb = doc.add_table(rows=len(risks), cols=1); set_no_border(t_rb); set_tw(t_rb)
        for i, risk in enumerate(risks):
            cell = t_rb.cell(i, 0); set_cell_bg(cell, AMBER_BG)
            set_cell_padding(cell, top=60, bottom=60, left=200, right=180)
            lines = risk.split('\n')
            ar(cp(cell), f'▶　{lines[0]}', size=9, color=AMBER, bold=True)
            if len(lines) > 1:
                p2 = cap(cell, sb=3); ar(p2, lines[1], size=8.5, color=DARK)
        np(doc, sa=5)
    colored_block(doc, '※ 上記はあくまでも「可能性」です',
        'すべてのお子さんに必ず起こるわけではありません。しかし、成長期に適切な介入を行うことで、これらのリスクを軽減できる可能性があります。今の状態を正しく理解し、選択肢を持っていただくことが大切です。',
        bg=LGRAY_BG, tc=DARK, bc=GRAY)
    np(doc, sa=4); doc.add_page_break()

    # ── PAGE 4：今 vs 後で 比較表 ──
    page_header(doc, f'{child_name}様　小児矯正診断レポート',
                '総括責任者：理事長　谷口正昭', dstr)
    section_label(doc, '今からできる治療　vs　永久歯列になってからの治療')
    t_key = doc.add_table(rows=1, cols=1); set_no_border(t_key); set_tw(t_key)
    c_key = t_key.cell(0, 0); set_cell_bg(c_key, NAVY)
    set_cell_padding(c_key, top=70, bottom=70, left=180, right=180)
    ar(cp(c_key), '小児矯正の目的は「単に歯並びを整えること」ではなく、「口元の機能を整えて、健康的な歯列を獲得すること」です', bold=True, size=9.5, color=WHITE)
    np(doc, sa=5)

    compare_rows = [
        ('主な目的',       '顎の成長・機能・スペース改善',                         '歯並び・噛み合わせのみ'),
        ('アプローチ',     '成長を利用する（今しかできない）',                      '歯を動かして整える'),
        ('対応できること', '口呼吸や異常嚥下癖など悪習癖の改善も可能',             '歯列の整列・細かい噛み合わせに対応'),
        ('抜歯',           '不要の可能性あり ✓',                                   '必要になる可能性が高い'),
        ('顔貌改善',       '横顔・骨格から改善できる ✓',                           '改善は難しい'),
        ('費用目安',       cost_1,                                                   '66〜99万円以上'),
        ('メリット',       '成長期の今だから機能改善を目指せる',                    '仕上がりを細かく整えられる'),
        ('注意点',         '装置使用・ご家庭での協力が必要',                        '抜歯・長期化の可能性'),
    ]
    ct = doc.add_table(rows=len(compare_rows)+1, cols=3)
    set_border(ct, color=LGRAY, size=4); set_tw(ct)
    for j, w in enumerate([3.5, 7.0, 7.5]): set_cw(ct, j, w)
    for j, (txt, bg, fc) in enumerate([('比較項目',LGRAY_BG,DARK),('小児矯正「ブランベラボキッズ」',NAVY,WHITE),('成人矯正',LGRAY,DARK)]):
        c = ct.cell(0, j); set_cell_bg(c, bg); set_cell_padding(c, top=70, bottom=70)
        p = cp(c); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(p, txt, bold=True, size=9, color=fc)
    for i, (label, now_txt, later_txt) in enumerate(compare_rows):
        row = ct.rows[i+1]
        c0 = row.cells[0]; set_cell_bg(c0, LGRAY_BG)
        set_cell_padding(c0, top=50, bottom=50, left=120)
        ar(cp(c0), label, bold=True, size=8.5, color=NAVY)
        c1 = row.cells[1]
        if i % 2 == 0: set_cell_bg(c1, LIGHT_BLUE)
        set_cell_padding(c1, top=50, bottom=50, left=120)
        ar(cp(c1), now_txt, size=8.5, color=GREEN if '✓' in now_txt else DARK)
        c2 = row.cells[2]
        if i % 2 == 0: set_cell_bg(c2, STRIPE)
        set_cell_padding(c2, top=50, bottom=50, left=120)
        ar(cp(c2), later_txt, size=8.5, color=AMBER if label in ('抜歯','顔貌改善') else DARK)
    np(doc, sa=4)

    t_rec2 = doc.add_table(rows=1, cols=1); set_no_border(t_rec2); set_tw(t_rec2)
    c_rec2 = t_rec2.cell(0, 0); set_cell_bg(c_rec2, GOLD_BG)
    set_cell_padding(c_rec2, top=80, bottom=80, left=180, right=180)
    ar(cp(c_rec2), 'Dr.推奨：今の成長期に小児矯正をスタートすることをおすすめします', bold=True, size=10, color=GOLD)
    p_rec2 = cap(c_rec2, sb=4)
    ar(p_rec2, f'推奨治療プラン：{rec_device}　｜　治療期間：{period_1}　｜　費用：{cost_1}', size=9, color=DARK)
    np(doc, sa=4); doc.add_page_break()

    # ── PAGE 5：理事長コメント ＋ ご家庭での協力 ──
    page_header(doc, f'{child_name}様　小児矯正診断レポート',
                '総括責任者：理事長　谷口正昭', dstr)
    section_label(doc, '担当医より　保護者の方へ')
    dc_t = doc.add_table(rows=1, cols=2); set_no_border(dc_t); set_tw(dc_t)
    set_cw(dc_t, 0, 0.3); set_cw(dc_t, 1, 17.7)
    set_cell_bg(dc_t.cell(0, 0), NAVY)
    cr_b = dc_t.cell(0, 1); set_cell_bg(cr_b, (0xF8, 0xF9, 0xFA))
    set_cell_padding(cr_b, top=90, bottom=90, left=220, right=180)
    ar(cp(cr_b), '理事長　谷口正昭より　保護者の方へ', bold=True, size=10, color=GOLD)
    p_dc = cap(cr_b, sb=6); p_dc.paragraph_format.space_after = Pt(10)
    ar(p_dc, f'「{comment}」', italic=True, size=9, color=DARK)
    p_ds = cap(cr_b); p_ds.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ar(p_ds, '医療法人 mirai　さいわいデンタルクリニック　理事長　谷口 正昭', bold=True, size=9, color=GOLD)
    np(doc, sa=6)

    section_label(doc, 'ご家庭でご協力いただきたいこと')
    t_home = doc.add_table(rows=1, cols=1); set_no_border(t_home); set_tw(t_home)
    c_home = t_home.cell(0, 0); set_cell_bg(c_home, LIGHT_BLUE)
    set_cell_padding(c_home, top=70, bottom=70, left=200, right=200)
    ar(cp(c_home), '小児矯正は、医院で装置を作るだけで完了する治療ではありません', bold=True, size=10, color=NAVY)
    p_hb = cap(c_home, sb=5)
    ar(p_hb, 'ご家庭での装置使用やトレーニングの継続が治療の成果に直結します。医院とご家庭が一緒に取り組むことで、より良い結果を目指します。', size=9, color=DARK)
    np(doc, sa=4)
    home_items = [
        ('装置の使用時間を守る', '決められた時間、装置を正しく使用してください。'),
        ('口を閉じる意識を持つ', 'お子さまが口を開けて過ごしていたら、優しく声をかけてあげてください。'),
        ('鼻呼吸を意識させる', '口呼吸から鼻呼吸への移行をサポートしてあげてください。'),
        ('トレーニングを継続する', '舌・唇のトレーニングを日課として続けていただけると効果が高まります。'),
        ('定期通院を継続する', '成長の確認と装置の調整のため、定期通院をお願いします。'),
    ]
    for title, detail in home_items:
        t_hi = doc.add_table(rows=1, cols=2); set_no_border(t_hi); set_tw(t_hi)
        set_cw(t_hi, 0, 4.5); set_cw(t_hi, 1, 13.5)
        c_hl = t_hi.cell(0, 0); set_cell_bg(c_hl, NAVY)
        set_cell_padding(c_hl, top=55, bottom=55, left=120, right=120); set_cell_valign(c_hl)
        ar(cp(c_hl), title, bold=True, size=8.5, color=WHITE)
        c_hr = t_hi.cell(0, 1); set_cell_bg(c_hr, STRIPE)
        set_cell_padding(c_hr, top=55, bottom=55, left=140, right=120); set_cell_valign(c_hr)
        ar(cp(c_hr), detail, size=8.5, color=DARK)
        np(doc, sa=1)
    np(doc, sa=4); doc.add_page_break()

    # ── PAGE 6：費用・期間・お支払い ──
    page_header(doc, f'{child_name}様　小児矯正診断レポート',
                '総括責任者：理事長　谷口正昭', dstr)
    section_label(doc, '費用・期間・お支払いについて')
    exp_lbl = expiry if expiry else '——'
    cost_rows = [
        ('治療名',              rec_device),
        ('小児矯正治療期間',    period_1),
        ('小児矯正治療費用',    cost_1),
        ('矯正治療管理料',      monthly_fee),
        ('装置再作製',          '追加費用なし（プラン内に含む）'),
        ('成人矯正',            phase2_note if phase2_note else '永久歯列完成後、必要に応じてご相談'),
        ('◎ デンタルクレジット', '12回払い　金利0円'),
        ('◎ 医療費控除',        '確定申告で一部還付可能（年間10万円超）'),
        ('→ 進める場合',        f'プログラム申込費用0円（期間限定・{exp_lbl}まで）'),
    ]
    ct_cost = doc.add_table(rows=len(cost_rows), cols=2)
    set_border(ct_cost, color=LGRAY, size=4); set_tw(ct_cost)
    set_cw(ct_cost, 0, 5.5); set_cw(ct_cost, 1, 12.5)
    for i, (lbl, cont) in enumerate(cost_rows):
        row = ct_cost.rows[i]; cl2 = row.cells[0]; cr2 = row.cells[1]
        if i % 2 == 0: set_cell_bg(cl2, LIGHT_BLUE); set_cell_bg(cr2, STRIPE)
        set_cell_padding(cl2, top=55, bottom=55, left=120)
        set_cell_padding(cr2, top=55, bottom=55, left=120)
        is_sp = lbl.startswith('◎') or lbl.startswith('→')
        ar(cp(cl2), lbl, bold=True, size=9, color=GOLD if is_sp else NAVY)
        ar(cp(cr2), cont, size=9, color=DARK)
    np(doc, sa=5)

    section_label(doc, '保護者の方からよくいただくご質問')

    def pedo_faq(q, a_lines):
        fq = doc.add_table(rows=2, cols=1); set_no_border(fq); set_tw(fq)
        cq = fq.cell(0, 0); set_cell_bg(cq, (0xF0, 0xF4, 0xF8))
        set_cell_padding(cq, top=45, bottom=45, left=160)
        ar(cp(cq), 'Q.　'+q, bold=True, size=9, color=NAVY)
        ca = fq.cell(1, 0)
        tcPr2 = ca._tc.get_or_add_tcPr()
        tcB = OxmlElement('w:tcBorders'); lb = OxmlElement('w:left')
        lb.set(qn('w:val'),'single'); lb.set(qn('w:sz'),'10')
        lb.set(qn('w:space'),'0'); lb.set(qn('w:color'),hex3(NAVY))
        tcB.append(lb); tcPr2.append(tcB)
        set_cell_padding(ca, top=45, bottom=55, left=180)
        for i, (line, bf, cf2) in enumerate(a_lines):
            p_a = cp(ca) if i == 0 else cap(ca, sb=3)
            ar(p_a, ('A.　' if i == 0 else '')+line, size=9, color=cf2, bold=bf)
        np(doc, sa=3)

    pedo_faq('小児矯正をすれば、将来また矯正しなくて済みますか？', [
        ('小児矯正は「口元の機能を整えて、健康的な歯列を獲得すること」が目的です。目標は小児矯正単独で歯列の完成を目指します。しかし、治療開始の年齢・歯並びや口腔機能の重症度によっては、永久歯が生え揃った後に、成人矯正が必要になる場合があります。', False, DARK),
        ('ただし、小児矯正を行うことで、将来の治療難易度が下がったり、抜歯が不要になる可能性が高まります。', False, DARK),
        ('当院では最初から成人矯正になった場合の費用のご案内や成人矯正へ移行時の割引プログラムがありますので、後から想定外の費用が発生することはありません。', True, NAVY),
    ])
    pedo_faq('何歳から始めるのが良いですか？', [
        ('お子さんの状態によりますが、顎の骨格が成長しているうちに始めるほど効果的です。', False, DARK),
        ('一般的に6〜10歳頃（混合歯列期）が、骨格へのアプローチが最もしやすい時期と言われています。', False, DARK),
        (f'{child_name}さんの場合は、今がまさにその時期にあたります。', True, NAVY),
    ])
    pedo_faq('装置を嫌がったらどうすればいいですか？', [
        ('はじめはお子さまが装置に慣れるまで時間がかかることがあります。', False, DARK),
        ('無理に強制するより、「なぜやるのか」をお子さまに合わせた言葉で伝えることが大切です。', False, DARK),
        ('当院では定期通院のたびにお子さまへの声かけもサポートします。困ったことがあればいつでもご連絡ください。', True, NAVY),
    ])
    np(doc, sa=4); doc.add_page_break()

    # ── PAGE 7：次のステップ ──
    page_header(doc, f'{child_name}様　小児矯正診断レポート',
                '総括責任者：理事長　谷口正昭', dstr)
    t_ns = doc.add_table(rows=1, cols=1); set_no_border(t_ns); set_tw(t_ns)
    cn = t_ns.cell(0, 0); set_cell_bg(cn, NAVY)
    set_cell_padding(cn, top=100, bottom=100, left=200, right=200)
    ar(cp(cn), '次のステップ', bold=True, size=11, color=WHITE)
    for title, detail in [
        ('① ご家族でこのレポートをご確認ください',
         '特に「今からできる治療 vs 永久歯列後の治療」と「費用・期間」のページをご覧ください。'),
        ('② 不安な点をLINEで送ってください',
         '装置の使用時間・痛み・費用・将来の成人矯正の必要性など、どんな内容でも大丈夫です。'),
        ('③ 必要であれば無料再相談が可能です',
         'ご両親おそろいでの相談、オンライン相談も可能です。「もう一度聞きたい」は大歓迎です。'),
        ('④ 治療を進める場合は装置作成・治療開始の手続きに進みます',
         '成長期の治療は開始時期も大切です。ご希望の場合は次回ご来院時に同意書・お支払い・装置作成の準備に進みます。'),
    ]:
        pt2 = cap(cn, sb=8); ar(pt2, title, bold=True, size=9.5, color=GOLD)
        pd2 = cap(cn, sb=2); ar(pd2, detail, size=8.5, color=WHITE)
    cap(cn, sb=10)
    for line in ['小児矯正は、今すぐ決める必要はありません。',
                 'ただし、成長は待ってくれません。',
                 '',
                 '今の時期にできることと、後からしかできないことがあります。',
                 'このレポートが、ご家族の判断の一助になれば幸いです。',
                 '',
                 '不安な点があれば、遠慮なくご相談ください。',
                 '納得して判断していただくことを、私たちは大切にしています。']:
        pc = cap(cn, sb=2); ar(pc, line, size=8.5, color=WHITE)
    np(doc, sa=5)

    t_cl = doc.add_table(rows=1, cols=1); set_tw(t_cl)
    tblPr_el = get_tblPr(t_cl._tbl); borders_el = OxmlElement('w:tblBorders')
    for b in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{b}'); el.set(qn('w:val'),'single')
        el.set(qn('w:sz'),'8'); el.set(qn('w:space'),'0')
        el.set(qn('w:color'),hex3(GOLD)); borders_el.append(el)
    for b in ['insideH','insideV']:
        el = OxmlElement(f'w:{b}'); el.set(qn('w:val'),'none'); borders_el.append(el)
    tblPr_el.append(borders_el)
    ccl = t_cl.cell(0, 0); set_cell_padding(ccl, top=80, bottom=80, left=150, right=150)
    pc1 = cp(ccl); pc1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _cn = d.get('clinic_name', CLINIC_LIST[0])
    _ci = d.get('clinic_info', CLINIC_DATA[CLINIC_LIST[0]])
    ar(pc1, _cn, bold=True, size=10, color=NAVY)
    _cl = [_ci['address'], f'TEL : {_ci["tel"]}']
    if _ci.get('hours'): _cl.append(_ci['hours'])
    for line in _cl:
        pc2 = cap(ccl, sb=3); pc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar(pc2, line, size=8.5, color=DARK)

    set_footer_logo(doc, logo_bytes)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


def convert_to_pdf(docx_bytes):
    try:
        from docx2pdf import convert
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            f.write(docx_bytes); tmp=f.name
        pdf=tmp.replace('.docx','.pdf')
        convert(tmp, pdf)
        data=open(pdf,'rb').read()
        os.unlink(tmp); os.unlink(pdf)
        return data
    except Exception:
        return None


# ════════════════════════════════════════════
# STREAMLIT UI
# ════════════════════════════════════════════

st.set_page_config(page_title='矯正診断レポート 生成ツール', page_icon='🦷', layout='wide')

st.markdown("""
<style>
h2 { color: #0D2B4E; }
.stButton > button {
    background-color: #0D2B4E; color: white;
    font-size: 16px; padding: 0.6em 2.5em; border-radius: 6px; border: none;
}
.stButton > button:hover { background-color: #1A4A7A; }
div[data-testid="stExpander"] { border: 1px solid #e0e0e0; border-radius: 6px; }
</style>
""", unsafe_allow_html=True)

# ── パスワード認証 ──
import streamlit as _st_auth
_APP_PASSWORD = st.secrets.get("APP_PASSWORD", "mirai2025") if hasattr(st, "secrets") else "mirai2025"
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if not st.session_state.authenticated:
    st.markdown("## 🦷 矯正診断レポート 自動生成ツール")
    st.markdown("#### 医療法人 mirai　スタッフ専用")
    st.markdown("---")
    _c1, _c2, _c3 = st.columns([1, 2, 1])
    with _c2:
        _pw = st.text_input("パスワード", type="password", placeholder="パスワードを入力してください")
        if st.button("ログイン", use_container_width=True):
            if _pw == _APP_PASSWORD:
                st.session_state.authenticated = True
                st.rerun()
            else:
                st.error("パスワードが違います")
    st.stop()

st.markdown('## 🦷　矯正診断レポート　自動生成ツール')
st.caption('さいわいデンタルクリニック　moyuk SAPPORO院　｜　医療法人 mirai')
st.markdown('---')

# ── 写真ヘルパー（両モードで使用） ──
def read_file(f): return f.read() if f else None

def _photo_selector(label, folder_key, upload_key):
    up = st.file_uploader(label, type=['jpg','jpeg','png'], key=upload_key)
    if up:
        b = read_file(up)
        if b: st.image(b, width=120)
        return b
    return None

# ── モード選択 ──
report_mode = st.radio(
    'レポートの種類を選択',
    ['成人矯正', '小児矯正'],
    horizontal=True,
    key='report_mode',
)
st.markdown('---')

# ════════════════════════════════════════════
# 小児矯正モード UI
# ════════════════════════════════════════════
if report_mode == '小児矯正':
    st.markdown('### 👶 基本情報（小児矯正）')
    pedo_clinic_name = st.selectbox('🏥 医院', CLINIC_LIST, key='pedo_clinic_select')
    p_col1, p_col2, p_col3, p_col4 = st.columns([2, 2, 1, 2])
    with p_col1:
        pedo_child_name = st.text_input('お子さんの名前 *', placeholder='例：山田 はなこ', key='pedo_child')
    with p_col2:
        pedo_guardian = st.text_input('保護者名 *', placeholder='例：山田 花子', key='pedo_guardian')
    with p_col3:
        pedo_age = st.text_input('年齢・学年', placeholder='例：8歳', key='pedo_age')
    with p_col4:
        pedo_date = st.date_input('診断日', value=datetime.date.today(), key='pedo_date')

    st.markdown('---')
    st.markdown('### 🔍 今日の検査でわかったこと（所見）')
    pedo_finding_cols = st.columns(2)
    pedo_selected_findings = []
    for _fi, _opt in enumerate(PEDO_FINDINGS_LIST):
        with pedo_finding_cols[_fi % 2]:
            if st.checkbox(_opt, key=f'pedo_finding_{_fi}'):
                pedo_selected_findings.append(_opt)
    pedo_space = st.text_input('スペース不足量（任意）', placeholder='例：約6mm', key='pedo_space')

    st.markdown('---')
    st.markdown('### 💊 推奨装置・治療計画')
    pd_c1, pd_c2, pd_c3 = st.columns(3)
    with pd_c1:
        pedo_device = st.text_input('推奨装置名', value='ブランベラボキッズプログラム', key='pedo_device')
    with pd_c2:
        pedo_period1 = st.text_input('一期治療期間', value='14歳ごろまで（成長による）', key='pedo_period1')
    with pd_c3:
        pedo_cost1 = st.text_input('一期治療費用', value='440,000円（税込）', key='pedo_cost1')
    pd_c4, pd_c5 = st.columns(2)
    with pd_c4:
        pedo_monthly = st.text_input('月次調整料', value='5,500円/月（税込）', key='pedo_monthly')
    with pd_c5:
        pedo_has_exp = st.checkbox('期間限定特典あり', value=True, key='pedo_has_exp')
        if pedo_has_exp:
            pedo_exp_date = st.date_input('有効期限', value=datetime.date.today() + datetime.timedelta(days=30), key='pedo_exp_date')
            pedo_exp_str = pedo_exp_date.strftime('%Y年%m月%d日')
        else:
            pedo_exp_str = ''
    pedo_phase2 = st.text_area(
        '二期治療について',
        value='永久歯列完成後、必要に応じてアライナー矯正による仕上げ治療を行う可能性があります。',
        height=70, key='pedo_phase2',
    )

    st.markdown('---')
    st.markdown('### ⚠️ 放置リスク（該当するものを選択）')
    pedo_risks = st.multiselect('リスク項目', PEDO_RISK_LIST, key='pedo_risks')

    st.markdown('---')
    st.markdown('### 💬 理事長コメント（保護者の方へ）')
    _PEDO_CMT_DEFAULT = (
        '今日の診断では、歯が並ぶためのスペース不足に加えて、噛み合わせや口呼吸など、'
        'お口の機能面にも改善が必要な状態が確認されました。\n\n'
        '小児矯正の目的は、単に今の前歯をきれいに並べることではありません。'
        '成長期の力を利用しながら、永久歯が並びやすい環境を整え、将来の噛み合わせや口元のバランスをより良い方向へ導くことが目的です。\n\n'
        '小児矯正を行ったからといって、将来の矯正治療が必ず不要になるわけではありません。'
        'しかし、今の時期に顎の幅・噛み合わせ・口呼吸・舌や唇の使い方にアプローチすることで、'
        '将来の治療の選択肢を大きく広げられる可能性があります。\n\n'
        'ご家族で十分にご相談いただき、不安な点は一つずつ確認しながら進めていきましょう。'
    )
    if 'pedo_comment_text' not in st.session_state:
        st.session_state['pedo_comment_text'] = _PEDO_CMT_DEFAULT
    pedo_comment = st.text_area('コメント', key='pedo_comment_text', height=220)

    st.markdown('---')
    st.markdown('### 📸 写真')
    ph_c1, ph_c2, ph_c3 = st.columns(3)
    with ph_c1:
        pedo_face_front = _photo_selector('顔貌（正面）', 'psel_ff', 'pup_ff')
        pedo_face_side  = _photo_selector('顔貌（横顔）', 'psel_fs', 'pup_fs')
    with ph_c2:
        pedo_oral  = _photo_selector('口腔内写真', 'psel_or', 'pup_or')
        pedo_scan  = _photo_selector('スキャンデータ', 'psel_sc', 'pup_sc')
    with ph_c3:
        pedo_dr    = _photo_selector('理事長の写真（任意）', 'psel_pdr', 'pup_pdr')

    st.markdown('---')
    p_btn_col, p_info_col = st.columns([1, 3])
    with p_btn_col:
        pedo_gen_btn = st.button('📄　小児矯正レポートを生成', use_container_width=True)
    with p_info_col:
        if not pedo_child_name:
            st.info('お子さんの名前を入力してから生成してください')

    if pedo_gen_btn:
        if not pedo_child_name:
            st.error('お子さんの名前を入力してください')
        else:
            with st.spinner('レポートを生成しています...'):
                pedo_data = {
                    'clinic_name':         pedo_clinic_name,
                    'clinic_info':         CLINIC_DATA[pedo_clinic_name],
                    'child_name':          pedo_child_name,
                    'guardian_name':       pedo_guardian,
                    'child_age':           pedo_age,
                    'date_str':            pedo_date.strftime('%Y年%m月%d日'),
                    'findings':            pedo_selected_findings,
                    'space_shortage':      pedo_space,
                    'recommended_device':  pedo_device,
                    'treatment_period_1':  pedo_period1,
                    'treatment_cost_1':    pedo_cost1,
                    'monthly_fee':         pedo_monthly,
                    'phase2_note':         pedo_phase2,
                    'risks_pedo':          pedo_risks,
                    'dr_comment':          pedo_comment,
                    'expiry_str':          pedo_exp_str,
                    'images': {
                        'face_front': pedo_face_front,
                        'face_side':  pedo_face_side,
                        'oral':       pedo_oral,
                        'scan':       pedo_scan,
                    },
                }
                docx_bytes_p = generate_pedo_report(pedo_data)

            st.success(f'✅ {pedo_child_name}様の小児矯正レポートが完成しました！')
            fname_p = f'小児矯正診断レポート_{pedo_child_name}_{pedo_date.strftime("%Y%m%d")}'
            dl1_p, dl2_p = st.columns(2)
            with dl1_p:
                st.download_button(
                    '📥 Word (.docx) をダウンロード',
                    data=docx_bytes_p,
                    file_name=f'{fname_p}.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    use_container_width=True,
                )
            with dl2_p:
                with st.spinner('PDF変換中...'):
                    pdf_bytes_p = convert_to_pdf(docx_bytes_p)
                if pdf_bytes_p:
                    st.download_button(
                        '📥 PDF をダウンロード',
                        data=pdf_bytes_p,
                        file_name=f'{fname_p}.pdf',
                        mime='application/pdf',
                        use_container_width=True,
                    )
                else:
                    st.info('PDF変換にはWordが必要です。Wordでdocxを開いてPDF保存してください。')
    st.stop()

# ════════════════════════════════════════════
# 成人矯正モード UI（既存）
# ════════════════════════════════════════════

# ── 基本情報 ──
st.markdown('### 📋 基本情報')
brand_option = st.radio(
    '🏷️ ブランド',
    ['さいわいデンタル', '行智会'],
    horizontal=True,
    key='brand_select',
)
_active_clinic_list = KOCHIKAI_CLINIC_LIST if brand_option == '行智会' else CLINIC_LIST
_active_clinic_data = KOCHIKAI_CLINIC_DATA if brand_option == '行智会' else CLINIC_DATA
clinic_name = st.selectbox('🏥 医院', _active_clinic_list, key='clinic_select')
col1, col2, col3 = st.columns([2, 1, 2])
with col1:
    patient_name = st.text_input('患者名 *', placeholder='例：山田 花子')
with col2:
    diag_date = st.date_input('診断日', value=datetime.date.today())
with col3:
    case_type = st.text_input('症例タイプ（Before/Afterのラベル）', placeholder='例：叢生、開咬、骨格性クラスII', value='叢生')

st.markdown('---')

# ── 患者の主訴（お悩み） ──
st.markdown('### 💭 患者さんの主訴・お悩み（初診時）')
st.caption('チェックした項目がレポート冒頭の「今回ご相談いただいたお悩み」ボックスに反映されます。')

_COMPLAINT_OPTIONS = [
    '前歯の歯並び（ガタつき）が気になる',
    '口元の出っ張り・突出感が気になる',
    '笑ったとき・写真のときの見た目を改善したい',
    '受け口（前歯が反対に噛んでいる）を治したい',
    '目立たない方法で治療したい',
    '噛み合わせや顎の違和感・痛みが気になる',
    '発音・滑舌が気になる',
    '将来のために虫歯・歯周病リスクを減らしたい',
    '歯を長く健康に保ちたい',
    '子どものうちに治してあげたい',
]
_complaint_cols = st.columns(2)
_selected_complaints = []
for _ci, _opt in enumerate(_COMPLAINT_OPTIONS):
    with _complaint_cols[_ci % 2]:
        if st.checkbox(_opt, key=f'complaint_{_ci}'):
            _selected_complaints.append(_opt)

complaint_free = st.text_input('その他・自由入力', placeholder='例：前歯が気になって人前で笑えない')
if complaint_free.strip():
    _selected_complaints.append(complaint_free.strip())

chief_complaints = _selected_complaints

st.markdown('---')

# ── 診断項目 ──
st.markdown('### 🔍 診断項目（最大4つ）')

# session_state init
for i in range(4):
    if f'desc_{i}' not in st.session_state: st.session_state[f'desc_{i}'] = ''
    if f'prev_key_{i}' not in st.session_state: st.session_state[f'prev_key_{i}'] = None

diagnoses = []
diag_cols = st.columns(2)
for i in range(4):
    with diag_cols[i % 2]:
        with st.expander(f'診断 {"①②③④"[i]}', expanded=(i < 2)):
            c_d, c_s = st.columns([3, 1])
            with c_d:
                diag_item = st.selectbox('診断項目', DIAG_LIST, key=f'diag_{i}')
            with c_s:
                severity = st.selectbox('重症度', SEV_LIST, key=f'sev_{i}')
            new_key = (diag_item, severity)
            if new_key != st.session_state[f'prev_key_{i}']:
                if diag_item != '（選択してください）':
                    st.session_state[f'desc_{i}'] = DIAG_DESC.get(new_key, f'{diag_item}（{severity}）の状態が認められます。')
                else:
                    st.session_state[f'desc_{i}'] = ''
                st.session_state[f'prev_key_{i}'] = new_key
            desc = st.text_area('診断説明文（編集可）', key=f'desc_{i}', height=80)
            diagnoses.append({'item': diag_item, 'severity': severity, 'description': desc})

st.markdown('---')

# ── リスク ──
st.markdown('### ⚠️ 放置リスク（該当するものを選択）')
selected_risks = st.multiselect('リスク項目', RISK_LIST)

st.markdown('---')

# ── 治療によって期待できる変化 ──
st.markdown('### ✅ 治療によって期待できる変化（該当するものを選択）')
BENEFIT_LIST = [
    '歯並びが整い、笑顔や口元の印象が自然に改善される可能性があります。',
    '歯磨きがしやすくなり、虫歯・歯周病リスクを長期的に下げられる可能性があります。',
    '咬み合わせのバランスが整い、特定の歯への過度な負担を軽減できる可能性があります。',
    '将来的な補綴治療（被せもの・インプラント等）や歯の破折リスクを下げられる可能性があります。',
    '口元・横顔の印象が自然に整う可能性があります。',
    '発音・滑舌が改善される可能性があります。',
    '咀嚼機能が向上し、食事・消化への負担が軽減される可能性があります。',
    '顎関節への負担が軽減され、顎の疲れや不快感が改善される可能性があります。',
    '自信を持って笑顔で人前に立てるようになる可能性があります。',
    '口腔内を清潔に保ちやすくなり、口臭予防にもつながる可能性があります。',
]
selected_benefits = st.multiselect('期待できる変化', BENEFIT_LIST)

st.markdown('---')

# ── 治療プラン ──
st.markdown('### 💊 治療プラン')
plan_names = list(PLAN_DATA.keys())

pc1, pc2, pc3 = st.columns(3)
with pc1:
    plan1 = st.selectbox('比較プラン①', plan_names, index=1)
with pc2:
    plan2 = st.selectbox('比較プラン②', plan_names, index=2)
with pc3:
    plan3 = st.selectbox('比較プラン③', plan_names, index=3)

selected_plans = [{'name': p} for p in [plan1, plan2, plan3]]
rec_options = [plan1, plan2, plan3]
rec_plan = st.selectbox('★ Dr.推奨プラン', rec_options, index=1)

col_pre, col_exp = st.columns([3, 1])
with col_pre:
    st.markdown('**矯正前に必要な処置**')
    pretx_clean   = st.checkbox('歯のクリーニング')
    pretx_cavity  = st.checkbox('虫歯治療')
    pretx_cavity_site = st.text_input('虫歯治療の部位', placeholder='例：上顎6番、下顎4番') if pretx_cavity else ''
    pretx_root    = st.checkbox('根の治療')
    pretx_root_site = st.text_input('根の治療の部位', placeholder='例：上顎7番') if pretx_root else ''
    pretx_wisdom  = st.checkbox('親知らずの抜歯')
    _items = []
    if pretx_clean:  _items.append('歯のクリーニング')
    if pretx_cavity: _items.append('虫歯治療' + (f'（{pretx_cavity_site}）' if pretx_cavity_site else ''))
    if pretx_root:   _items.append('根の治療' + (f'（{pretx_root_site}）' if pretx_root_site else ''))
    if pretx_wisdom: _items.append('親知らずの抜歯')
    pretreatment = '\n'.join(_items) if _items else '特になし'
with col_exp:
    has_expiry = st.checkbox('期間限定特典あり', value=True)
    if has_expiry:
        expiry_date = st.date_input('有効期限', value=datetime.date.today() + datetime.timedelta(days=30))
        expiry_str = expiry_date.strftime('%Y年%m月%d日')
    else:
        expiry_str = ''

st.markdown('---')

# ── 結論ボックス（1ページ目冒頭） ──
st.markdown('### 📋 結論ボックス（1ページ目の最上部に表示）')
st.caption('患者・ご家族が最初に目にする「あなた専用の結論」です。推奨プランが変わると自動で更新されます。')

_pname_c = patient_name if patient_name else '〇〇'
_first_diag_c = st.session_state.get('diag_0', '（選択してください）')
_first_diag_c = _first_diag_c if _first_diag_c != '（選択してください）' else '歯並び・咬み合わせの問題'
_rec_c = rec_plan if rec_plan else '（未選択）'
_conclusion_auto = (
    f'{_pname_c}様の場合、{_first_diag_c}が認められます。\n\n'
    f'見た目の改善だけであれば部分的な治療も選択肢になりますが、長期的な安定性や噛み合わせまで考えると、全体のバランスを整える治療が望ましい状態です。\n\n'
    f'そのため、当院では\n推奨プラン：{_rec_c}\nをおすすめします。'
)
_conclusion_key = (_pname_c, _first_diag_c, _rec_c)
if 'conclusion_text' not in st.session_state:
    st.session_state['conclusion_text'] = _conclusion_auto
if st.session_state.get('_last_conclusion_key') != _conclusion_key:
    st.session_state['conclusion_text'] = _conclusion_auto
    st.session_state['_last_conclusion_key'] = _conclusion_key

conclusion = st.text_area('結論テキスト（直接編集できます）', key='conclusion_text', height=160)

st.markdown('---')

# ── Dr.コメント（推奨プランの根拠） ──
st.markdown('### 💬 理事長　谷口正昭より　推奨プランの根拠コメント')

_pname_for_comment = patient_name if patient_name else '〇〇'
_plan_labels = [f'{p["label"]}：{p["title"]}' for p in PLAN_COMMENT_PRESETS]

selected_plan_idx = st.radio(
    'コメントパターンを選択',
    range(len(_plan_labels)),
    format_func=lambda i: _plan_labels[i],
    key='plan_comment_idx'
)
_auto_comment = PLAN_COMMENT_PRESETS[selected_plan_idx]['body'].replace('〇〇', _pname_for_comment)

if 'dr_comment_text' not in st.session_state:
    st.session_state['dr_comment_text'] = _auto_comment
if st.session_state.get('_last_plan_preset') != (selected_plan_idx, _pname_for_comment):
    st.session_state['dr_comment_text'] = _auto_comment
    st.session_state['_last_plan_preset'] = (selected_plan_idx, _pname_for_comment)

dr_comment = st.text_area('コメント（Word上でも直接編集できます）',
                           key='dr_comment_text', height=220)

st.markdown('---')

# ── シミュレーションURL ──
st.markdown('### 🔗 シミュレーションURL（QRコード自動生成）')
sim_url = st.text_input('シミュレーションURL', placeholder='https://...')
if sim_url:
    st.caption('✓ QRコードを自動生成してレポートに挿入します')

st.markdown('---')

# ── 画像アップロード ──
st.markdown('### 📸 画像')

_PHOTO_EXT = ('.jpg','.jpeg','.png','.JPG','.JPEG','.PNG')
def _load_case(fname):
    if fname == '（未使用）': return None
    p = os.path.join(CASES_FOLDER, fname)
    return open(p,'rb').read() if os.path.exists(p) else None

def _photo_selector(label, folder_key, upload_key):
    """アップロード形式。bytesを返す。"""
    up = st.file_uploader(label, type=['jpg','jpeg','png'], key=upload_key)
    if up:
        b = read_file(up)
        if b: st.image(b, width=120)
        return b
    return None

_ORAL_COMMENTS = [
    '（コメントなし）',
    '前歯の歯並びに乱れ（叢生）があり、歯ブラシが届きにくい部分があります。',
    '前歯が深く噛み込んでいます（過蓋咬合）。下の歯肉への負担も確認されます。',
    '前歯が噛み合っていない状態（開咬）です。奥歯への負担が増加しています。',
    '上下の前歯の咬み合わせが逆（反対咬合）になっています。',
    '前歯に隙間が見られます（空隙歯列）。歯のバランス改善が必要な状態です。',
    '上下の歯の中心線（正中）にズレがあります。見た目と咬み合わせ両面に関係します。',
    '歯の重なりが強く、スペース不足の状態です。',
]
_SCAN_COMMENTS = [
    '（コメントなし）',
    '3Dスキャンデータで歯列全体のバランスを確認しました。矯正治療の計画に使用します。',
    '奥歯の咬み合わせにズレがあり、顎関節や歯への負担が確認されます。',
    '歯列のアーチ形態に問題があり、スペース不足が生じています。',
    '上下の咬合平面にズレがあり、左右で噛み合わせのバランスが異なっています。',
    'スキャンデータにより、治療前の現状を正確に記録しています。',
]

img_col1, img_col2, img_col3 = st.columns(3)
with img_col1:
    st.markdown('**口腔内写真・スキャンデータ（P.1）**')
    oral_photo_bytes = _photo_selector('口腔内写真', 'sel_oral', 'up_oral')
    oral_photo_comment = st.selectbox('口腔内写真のコメント', _ORAL_COMMENTS, key='oral_comment')
    if oral_photo_comment == '（コメントなし）': oral_photo_comment = ''
    oral_comment_free = st.text_input('口腔内写真　自由入力', key='oral_comment_free', placeholder='独自コメントはこちらへ')
    if oral_comment_free.strip(): oral_photo_comment = oral_comment_free.strip()

    scan_data_bytes  = _photo_selector('スキャンデータ', 'sel_scan', 'up_scan')
    scan_data_comment = st.selectbox('スキャンデータのコメント', _SCAN_COMMENTS, key='scan_comment')
    if scan_data_comment == '（コメントなし）': scan_data_comment = ''
    scan_comment_free = st.text_input('スキャンデータ　自由入力', key='scan_comment_free', placeholder='独自コメントはこちらへ')
    if scan_comment_free.strip(): scan_data_comment = scan_comment_free.strip()
with img_col2:
    st.markdown(f'**Before / After（P.3）**')
    st.caption(f'症例フォルダ: cases/　← 写真を入れると選択できます')
    before1_bytes = _photo_selector('Before ①', 'sel_b1', 'up_b1')
    after1_bytes  = _photo_selector('After ①',  'sel_a1', 'up_a1')
    before2_bytes = _photo_selector('Before ②', 'sel_b2', 'up_b2')
    after2_bytes  = _photo_selector('After ②',  'sel_a2', 'up_a2')
with img_col3:
    st.markdown('**理事長写真（P.3）**')
    dr_photo_bytes = _photo_selector('理事長の写真', 'sel_dr', 'up_dr')

st.markdown('---')

# ── 生成ボタン ──
col_btn, col_info = st.columns([1, 3])
with col_btn:
    generate_btn = st.button('📄　診断レポートを生成', use_container_width=True)
with col_info:
    if not patient_name:
        st.info('患者名を入力してから生成してください')

if generate_btn:
    if not patient_name:
        st.error('患者名を入力してください')
    else:
        with st.spinner('レポートを生成しています...'):
            data = {
                'brand': 'kochikai' if brand_option == '行智会' else 'saiwai',
                'clinic_name': clinic_name,
                'clinic_info': _active_clinic_data[clinic_name],
                'patient_name': patient_name,
                'date_str': diag_date.strftime('%Y年%m月%d日'),
                'diagnoses': diagnoses,
                'risks': selected_risks,
                'benefits': selected_benefits,
                'plans': selected_plans,
                'recommended_plan': rec_plan,
                'pretreatment': pretreatment,
                'chief_complaints': chief_complaints,
                'conclusion': conclusion,
                'dr_comment': dr_comment,
                'expiry_str': expiry_str,
                'case_type': case_type,
                'simulation_url': sim_url,
                'photo_comments': {
                    'oral_photo': oral_photo_comment,
                    'scan_data':  scan_data_comment,
                },
                'images': {
                    'oral_photo': oral_photo_bytes,
                    'scan_data':  scan_data_bytes,
                    'before1':    before1_bytes,
                    'after1':     after1_bytes,
                    'before2':    before2_bytes,
                    'after2':     after2_bytes,
                    'dr_photo':   dr_photo_bytes,
                },
            }
            docx_bytes = generate_report(data)

        st.success(f'✅ {patient_name}様のレポートが完成しました！')
        fname = f'矯正診断レポート_{patient_name}_{diag_date.strftime("%Y%m%d")}'

        dl1, dl2 = st.columns(2)
        with dl1:
            st.download_button(
                '📥 Word (.docx) をダウンロード',
                data=docx_bytes,
                file_name=f'{fname}.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                use_container_width=True,
            )
        with dl2:
            with st.spinner('PDF変換中...'):
                pdf_bytes = convert_to_pdf(docx_bytes)
            if pdf_bytes:
                st.download_button(
                    '📥 PDF をダウンロード',
                    data=pdf_bytes,
                    file_name=f'{fname}.pdf',
                    mime='application/pdf',
                    use_container_width=True,
                )
            else:
                st.info('PDF変換にはWordが必要です。Wordでdocxを開いてPDF保存してください。')
